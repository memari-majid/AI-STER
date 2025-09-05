#!/usr/bin/env python3
"""
AI-STER: AI-Powered Student Teaching Evaluation System

Copyright © 2025 Utah Valley University School of Education
All Rights Reserved.

This software is proprietary and confidential property of Utah Valley University
School of Education. Licensed for educational use only within accredited 
educational institutions for student teacher evaluation and assessment.

For licensing inquiries: email school_of_education@uvu.edu

IMPORTANT: This software is protected by copyright law. Unauthorized copying,
distribution, or modification is strictly prohibited.
"""

import streamlit as st
import pandas as pd
import json
from datetime import datetime, date
import uuid
from typing import Dict, List, Optional
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Import our modules
from data.rubrics import get_field_evaluation_items, get_ster_items, get_professional_dispositions, filter_items_by_evaluator_role
from data.synthetic import generate_synthetic_evaluations
from services.openai_service import OpenAIService
from services.pdf_service import PDFService
from utils.storage import save_evaluation, load_evaluations, export_data, import_data, save_ai_original, get_evaluation_comparison, get_evaluation_by_id
from utils.validation import validate_evaluation, calculate_score

# Page configuration
st.set_page_config(
    page_title="AI-STER - Student Teaching Evaluation System",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS following UVU brand guidelines - Primary colors lead, minimal accent use
st.markdown("""
<style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* UVU Brand Colors - Following official guidelines */
    :root {
        --uvu-green: #185C33;        /* Primary - should dominate */
        --white: #FFFFFF;            /* Primary - for clarity and space */
        --wolverine-green: #008A40;  /* Secondary - use sparingly */
        --grey: #AAAAAB;             /* Secondary - subtle support */
        --warm: #F2F0EB;             /* Secondary - light backgrounds */
        --black: #000000;            /* Secondary - text */
        /* Accent colors - use minimally (< 20% of design) */
        --valley-rise: #78BE3F;
        --lake-calm: #87C7BA;
        --heritage-brick: #B45336;
        --legacy-gold: #D2AC5F;
    }
    
    /* Global Styles - Clean and professional */
    .stApp {
        font-family: 'Inter', sans-serif;
        background-color: var(--white);
    }
    
    /* Headers - Simple and professional */
    h1, h2, h3 {
        font-family: 'Rajdhani', sans-serif !important;
        color: var(--uvu-green) !important;
        font-weight: 600 !important;
    }
    
    h1 {
        font-size: 2.5rem !important;
        margin-bottom: 1rem !important;
        border: none !important;
    }
    
    /* Sidebar - Clean with subtle background */
    section[data-testid="stSidebar"] {
        background-color: var(--warm) !important;
    }
    
    /* Primary buttons - UVU green with clean design */
    .stButton > button {
        background-color: var(--uvu-green);
        color: var(--white);
        border: none;
        border-radius: 6px;
        padding: 0.5rem 1.5rem;
        font-weight: 500;
        font-size: 1rem;
        transition: background-color 0.2s ease;
    }
    
    .stButton > button:hover {
        background-color: #0f3d1f;  /* Darker shade of UVU green */
    }
    
    /* Download buttons - Subtle secondary style */
    .stDownloadButton > button {
        background-color: var(--white);
        color: var(--uvu-green);
        border: 2px solid var(--uvu-green);
        border-radius: 6px;
        font-weight: 500;
    }
    
    .stDownloadButton > button:hover {
        background-color: var(--warm);
    }
    
    /* Alert messages - Subtle with minimal color */
    .stSuccess {
        background-color: rgba(24, 92, 51, 0.05);
        border-left: 3px solid var(--uvu-green);
        color: var(--black);
    }
    
    .stInfo {
        background-color: var(--warm);
        border-left: 3px solid var(--grey);
        color: var(--black);
    }
    
    .stWarning {
        background-color: rgba(210, 172, 95, 0.1);
        border-left: 3px solid var(--legacy-gold);
        color: var(--black);
    }
    
    .stError {
        background-color: rgba(180, 83, 54, 0.1);
        border-left: 3px solid var(--heritage-brick);
        color: var(--black);
    }
    
    /* Metrics - Clean and minimal */
    [data-testid="metric-container"] {
        background-color: var(--white);
        border: 1px solid #e0e0e0;
        padding: 1rem;
        border-radius: 6px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    
    [data-testid="metric-container"] [data-testid="metric-label"] {
        color: var(--grey);
        font-weight: 500;
        font-size: 0.875rem;
    }
    
    [data-testid="metric-container"] [data-testid="metric-value"] {
        color: var(--uvu-green);
        font-weight: 600;
    }
    
    /* Form inputs - Clean borders */
    .stSelectbox > div > div,
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        border: 1px solid #d0d0d0;
        border-radius: 4px;
    }
    
    .stSelectbox > div > div:hover,
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: var(--uvu-green);
        outline: none;
        box-shadow: 0 0 0 1px var(--uvu-green);
    }
    
    /* Radio buttons - Subtle green */
    .stRadio > div > label > div:first-child > div {
        background-color: var(--uvu-green);
    }
    
    /* Expanders - Clean design */
    .streamlit-expanderHeader {
        background-color: var(--white);
        border: 1px solid #e0e0e0;
        border-radius: 4px;
        font-weight: 500;
        color: var(--black);
    }
    
    .streamlit-expanderHeader:hover {
        background-color: var(--warm);
    }
    
    /* DataFrames - Clean borders */
    .stDataFrame {
        border: 1px solid #e0e0e0;
        border-radius: 4px;
    }
    
    /* Tabs - Professional look */
    .stTabs [data-baseweb="tab-list"] {
        background-color: var(--white);
        border-bottom: 2px solid #e0e0e0;
    }
    
    .stTabs [data-baseweb="tab"] {
        color: var(--grey);
        font-weight: 500;
        background-color: transparent;
        border-bottom: 2px solid transparent;
        margin-bottom: -2px;
    }
    
    .stTabs [aria-selected="true"] {
        color: var(--uvu-green) !important;
        background-color: transparent !important;
        border-bottom: 2px solid var(--uvu-green) !important;
    }
    
    /* Dividers - Subtle */
    hr {
        border: none;
        border-top: 1px solid #e0e0e0;
        margin: 2rem 0;
    }
    
    /* File uploader - Clean dashed border */
    .stFileUploader > div {
        border: 2px dashed #d0d0d0;
        border-radius: 4px;
        background-color: var(--white);
        transition: all 0.2s ease;
    }
    
    .stFileUploader > div:hover {
        border-color: var(--uvu-green);
        background-color: var(--warm);
    }
    
    /* Progress bars - Subtle green */
    .stProgress > div > div > div {
        background-color: var(--uvu-green);
    }
    
    /* Remove excessive shadows and animations */
    .element-container {
        transition: none;
    }
    
    /* Ensure white space and breathing room */
    .main .block-container {
        padding: 2rem 1rem;
        max-width: 1200px;
        margin: 0 auto;
    }
    
    /* Clean, professional link styling */
    a {
        color: var(--uvu-green);
        text-decoration: none;
    }
    
    a:hover {
        text-decoration: underline;
    }
</style>
""", unsafe_allow_html=True)

# Initialize services
openai_service = OpenAIService()
pdf_service = PDFService()

def main():
    """Main application entry point"""
    
    # Clean, professional header following UVU guidelines
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("""
            <h1 style="color: #185C33; margin: 0;">🎓 AI-STER</h1>
            <p style="color: #666666; margin: 0; font-size: 1.1rem;">
                Student Teaching Evaluation Rubric System • Utah Valley University
            </p>
        """, unsafe_allow_html=True)
    
    with col2:
        # Simple status indicator
        if openai_service.is_enabled():
            st.success("AI Features Enabled")
        else:
            st.info("Configure AI in Settings")
    
    # Sidebar navigation - clean and professional
    with st.sidebar:
        # Logo
        st.image("logo.png", width=180)
        
        st.markdown("### Navigation")
        
        page = st.selectbox(
            "",
            ["📝 New Evaluation", "📊 Dashboard", "🧪 Test Data", "⚙️ Settings"],
            index=0,
            key="navigation",
            label_visibility="collapsed"
        )
        
        st.markdown("---")
        
        # Quick statistics
        st.markdown("### Quick Statistics")
        
        evaluations = load_evaluations()
        
        # Simple metrics display
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Total", len(evaluations))
            st.metric("Drafts", len([e for e in evaluations if e.get('status') == 'draft']))
        with col2:
            st.metric("Completed", len([e for e in evaluations if e.get('status') == 'completed']))
            if evaluations:
                lesson_plan_count = len([e for e in evaluations if e.get('lesson_plan_provided', False)])
                submission_rate = (lesson_plan_count / len(evaluations)) * 100
                st.metric("LP Rate", f"{submission_rate:.0f}%")
    
    # Route to different pages
    if page == "📊 Dashboard":
        show_dashboard()
    elif page == "📝 New Evaluation":
        show_evaluation_form()
    elif page == "🧪 Test Data":
        show_test_data()
    elif page == "⚙️ Settings":
        show_settings()

def show_dashboard():
    """Dashboard with evaluation overview and analytics"""
    st.header("📊 Evaluation Dashboard")
    st.markdown("Track your evaluation progress and analytics")
    
    evaluations = load_evaluations()
    
    if not evaluations:
        st.info("No evaluations found. Create your first evaluation to get started!")
        if st.button("Create New Evaluation"):
            st.session_state.page = "📝 New Evaluation"
            st.rerun()
        return
    
    # Convert to DataFrame for analysis
    df = pd.DataFrame(evaluations)
    
    # Enhanced Metrics Row 1
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Evaluations", len(df))
    with col2:
        if 'status' in df.columns:
            completed = len(df[df['status'] == 'completed'])
        else:
            completed = 0
        st.metric("Completed", completed)
    with col3:
        if completed > 0 and 'total_score' in df.columns:
            avg_score = df[df['status'] == 'completed']['total_score'].mean()
            st.metric("Average Score", f"{avg_score:.1f}")
        else:
            st.metric("Average Score", "N/A")
    with col4:
        if len(df) > 0 and 'status' in df.columns:
            success_rate = (completed/len(df)*100)
            st.metric("Success Rate", f"{success_rate:.1f}%")
        else:
            st.metric("Success Rate", "N/A")
    
    # Enhanced Metrics Row 2 - Dashboard Feedback Implementation
    st.markdown("---")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        # Evaluation Types breakdown
        field_evals = len(df[df['rubric_type'] == 'field_evaluation']) if 'rubric_type' in df.columns else 0
        ster_evals = len(df[df['rubric_type'] == 'ster']) if 'rubric_type' in df.columns else 0
        st.metric("Field Evaluations", field_evals)
        st.metric("STER Evaluations", ster_evals)
    with col2:
        # Subject Areas count
        unique_subjects = df['subject_area'].nunique() if 'subject_area' in df.columns else 0
        total_subjects = len(df['subject_area'].dropna()) if 'subject_area' in df.columns else 0
        st.metric("Subject Areas", unique_subjects)
        if total_subjects > 0:
            st.caption(f"Total subject records: {total_subjects}")
    with col3:
        # Current Semester
        current_semester = df['semester'].mode()[0] if 'semester' in df.columns and len(df) > 0 else "Spring 2025"
        semester_count = len(df[df['semester'] == current_semester]) if 'semester' in df.columns else 0
        st.metric("Current Semester", current_semester)
        st.caption(f"Evaluations this semester: {semester_count}")
    with col4:
        # Department distribution
        unique_departments = df['department'].nunique() if 'department' in df.columns else 0
        st.metric("Departments", unique_departments)
        if 'department' in df.columns:
            dept_counts = df['department'].value_counts()
            most_active = dept_counts.index[0] if len(dept_counts) > 0 else "N/A"
            st.caption(f"Most active: {most_active}")
    
    # Enhanced Charts Section
    st.subheader("📈 Evaluation Analytics")
    
    # Row 1: Status and Type Charts (existing)
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Evaluation Status")
        if 'status' in df.columns:
            status_counts = df['status'].value_counts()
            st.bar_chart(status_counts)
        else:
            st.info("No status data available")
    
    with col2:
        st.subheader("Evaluations by Type")
        if 'rubric_type' in df.columns:
            type_counts = df['rubric_type'].value_counts()
            st.bar_chart(type_counts)
        else:
            st.info("No type data available")
    
    # Row 2: New Enhanced Charts
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Score Distribution")
        if 'total_score' in df.columns and completed > 0:
            completed_df = df[df['status'] == 'completed']
            score_bins = pd.cut(completed_df['total_score'], bins=[0, 10, 15, 20, 25, 30], labels=['0-10', '11-15', '16-20', '21-25', '26+'])
            score_dist = score_bins.value_counts().sort_index()
            st.bar_chart(score_dist)
        else:
            st.info("No score data available")
    
    with col2:
        st.subheader("Subject Areas")
        if 'subject_area' in df.columns:
            subject_counts = df['subject_area'].value_counts()
            st.bar_chart(subject_counts)
        else:
            st.info("No subject data available")
    
    # Row 3: School Context Charts
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("School Settings")
        if 'school_setting' in df.columns:
            setting_counts = df['school_setting'].value_counts()
            st.bar_chart(setting_counts)
        else:
            st.info("No school setting data available")
    
    with col2:
        st.subheader("Departments")
        if 'department' in df.columns:
            dept_counts = df['department'].value_counts()
            st.bar_chart(dept_counts)
            
            # Show department breakdown
            with st.expander("Department Details"):
                for dept, count in dept_counts.items():
                    percentage = (count / len(df)) * 100
                    st.write(f"**{dept}**: {count} evaluations ({percentage:.1f}%)")
        else:
            st.info("No department data available")
    
    # Enhanced Competency Area Performance Analysis
    st.subheader("🎯 Competency Area Performance")
    
    completed_evals = [e for e in evaluations if e.get('status') == 'completed']
    if completed_evals:
        competency_analysis = analyze_competency_performance(completed_evals)
        
        if competency_analysis:
            # Create competency performance chart
            comp_df = pd.DataFrame(list(competency_analysis.items()), columns=['Competency Area', 'Average Score'])
            comp_df = comp_df.sort_values('Average Score', ascending=True)
            
            st.bar_chart(comp_df.set_index('Competency Area'))
            
            # Enhanced performance metrics
            col1, col2, col3 = st.columns(3)
            
            strong_areas = [area for area, score in competency_analysis.items() if score >= 2.5]
            meeting_areas = [area for area, score in competency_analysis.items() if 2.0 <= score < 2.5]
            growth_areas = [area for area, score in competency_analysis.items() if score < 2.0]
            
            with col1:
                st.metric("🟢 Strong Performance", len(strong_areas), help="Areas scoring 2.5+ on average")
                if strong_areas:
                    st.caption("Top areas: " + ", ".join(strong_areas[:2]))
            
            with col2:
                st.metric("🟡 Meeting Standards", len(meeting_areas), help="Areas scoring 2.0-2.4 on average")
                if meeting_areas:
                    st.caption("Steady areas: " + ", ".join(meeting_areas[:2]))
            
            with col3:
                st.metric("🔴 Needs Attention", len(growth_areas), delta=f"{-len(growth_areas) if growth_areas else 0}", delta_color="inverse", help="Areas scoring below 2.0")
                if growth_areas:
                    st.caption("Focus areas: " + ", ".join(growth_areas[:2]))
            
            # Enhanced detailed breakdown
            with st.expander("📊 Detailed Competency Analysis", expanded=False):
                st.markdown("**Performance by Evaluation Type:**")
                
                # Analyze by rubric type
                field_evals = [e for e in completed_evals if e.get('rubric_type') == 'field_evaluation']
                ster_evals = [e for e in completed_evals if e.get('rubric_type') == 'ster']
                
                col1, col2 = st.columns(2)
                
                with col1:
                    if field_evals:
                        st.markdown("**Field Evaluations:**")
                        field_analysis = analyze_competency_performance(field_evals)
                        for area, avg_score in sorted(field_analysis.items(), key=lambda x: x[1], reverse=True):
                            if avg_score >= 2.5:
                                st.success(f"{area}: {avg_score:.2f}")
                            elif avg_score >= 2.0:
                                st.warning(f"{area}: {avg_score:.2f}")
                            else:
                                st.error(f"{area}: {avg_score:.2f}")
                    else:
                        st.info("No field evaluation data")
                
                with col2:
                    if ster_evals:
                        st.markdown("**STER Evaluations:**")
                        ster_analysis = analyze_competency_performance(ster_evals)
                        for area, avg_score in sorted(ster_analysis.items(), key=lambda x: x[1], reverse=True):
                            if avg_score >= 2.5:
                                st.success(f"{area}: {avg_score:.2f}")
                            elif avg_score >= 2.0:
                                st.warning(f"{area}: {avg_score:.2f}")
                            else:
                                st.error(f"{area}: {avg_score:.2f}")
                    else:
                        st.info("No STER evaluation data")
                
                # Department analysis
                if 'department' in df.columns:
                    st.markdown("---")
                    st.markdown("**Performance by Department:**")
                    
                    departments = df['department'].unique()
                    for dept in departments:
                        dept_evals = [e for e in completed_evals if e.get('department') == dept]
                        if dept_evals:
                            dept_analysis = analyze_competency_performance(dept_evals)
                            if dept_analysis:
                                avg_performance = sum(dept_analysis.values()) / len(dept_analysis)
                                st.write(f"**{dept}**: {len(dept_evals)} evaluations, avg performance: {avg_performance:.2f}")
    else:
        st.info("No completed evaluations for competency analysis")
    
    # Enhanced Professional Dispositions Summary (Field Evaluations Only)
    field_evals = [e for e in completed_evals if e.get('rubric_type') == 'field_evaluation']
    if field_evals:
        st.subheader("🌟 Professional Dispositions Summary")
        st.caption("Analysis of professional dispositions from field evaluations only")
        
        disposition_analysis = analyze_disposition_performance(field_evals)
        
        if disposition_analysis:
            disp_df = pd.DataFrame(list(disposition_analysis.items()), columns=['Disposition', 'Average Score'])
            disp_df = disp_df.sort_values('Average Score', ascending=True)
            
            st.bar_chart(disp_df.set_index('Disposition'))
            
            # Enhanced disposition metrics
            col1, col2, col3, col4 = st.columns(4)
            
            excellent_dispositions = [disp for disp, score in disposition_analysis.items() if score >= 3.5]
            meeting_dispositions = [disp for disp, score in disposition_analysis.items() if 3.0 <= score < 3.5]
            concerning_dispositions = [disp for disp, score in disposition_analysis.items() if score < 3.0]
            total_evaluations = len(field_evals)
            
            with col1:
                st.metric("Field Evaluations", total_evaluations, help="Total field evaluations with disposition data")
                
            with col2:
                st.metric("🟢 Excellent (3.5+)", len(excellent_dispositions), help="Dispositions scoring 3.5+ on average")
                
            with col3:
                st.metric("🟡 Meeting (3.0+)", len(meeting_dispositions), help="Dispositions meeting 3.0+ requirement")
                
            with col4:
                st.metric("🔴 Below Standard", len(concerning_dispositions), 
                         delta=f"{-len(concerning_dispositions) if concerning_dispositions else 0}", 
                         delta_color="inverse",
                         help="Dispositions below 3.0 requirement")
            
            # Critical alerts and detailed breakdown
            if concerning_dispositions:
                st.error(f"⚠️ **Critical Alert**: {len(concerning_dispositions)} dispositions below Level 3 requirement")
                st.markdown("**Dispositions requiring immediate attention:**")
                for disp in concerning_dispositions:
                    score = disposition_analysis[disp]
                    st.write(f"🔴 **{disp}**: {score:.2f} (needs to reach 3.0+)")
                
                st.markdown("**📋 Action Items:**")
                st.write("• Review student teacher performance in concerning disposition areas")
                st.write("• Provide targeted professional development and support")
                st.write("• Schedule additional observations focused on these dispositions")
                st.write("• Document improvement plans and progress monitoring")
            else:
                st.success("✅ **All dispositions meeting requirements** (Level 3+)")
                st.markdown("🎉 Strong professional disposition performance across all areas!")
            
            # Enhanced detailed analysis
            with st.expander("📊 Detailed Disposition Analysis", expanded=False):
                st.markdown("**Performance by Department:**")
                
                if 'department' in df.columns:
                    departments = df['department'].unique()
                    for dept in departments:
                        dept_field_evals = [e for e in field_evals if e.get('department') == dept]
                        if dept_field_evals:
                            dept_disp_analysis = analyze_disposition_performance(dept_field_evals)
                            if dept_disp_analysis:
                                avg_dept_disp = sum(dept_disp_analysis.values()) / len(dept_disp_analysis)
                                below_standard = [d for d, s in dept_disp_analysis.items() if s < 3.0]
                                
                                col1, col2 = st.columns([2, 1])
                                with col1:
                                    st.write(f"**{dept}** ({len(dept_field_evals)} evaluations)")
                                    if below_standard:
                                        st.write(f"⚠️ {len(below_standard)} disposition(s) below standard")
                                with col2:
                                    if avg_dept_disp >= 3.5:
                                        st.success(f"Avg: {avg_dept_disp:.2f}")
                                    elif avg_dept_disp >= 3.0:
                                        st.warning(f"Avg: {avg_dept_disp:.2f}")
                                    else:
                                        st.error(f"Avg: {avg_dept_disp:.2f}")
                
                st.markdown("---")
                st.markdown("**Individual Disposition Breakdown:**")
                for disp, score in sorted(disposition_analysis.items(), key=lambda x: x[1], reverse=True):
                    col1, col2, col3 = st.columns([3, 1, 1])
                    with col1:
                        st.write(f"**{disp}**")
                    with col2:
                        if score >= 3.5:
                            st.success(f"{score:.2f}")
                        elif score >= 3.0:
                            st.warning(f"{score:.2f}")
                        else:
                            st.error(f"{score:.2f}")
                    with col3:
                        # Calculate how many students are below standard for this disposition
                        below_count = sum(1 for eval in field_evals 
                                        if eval.get('disposition_scores', {}).get(disp, 0) < 3)
                        if below_count > 0:
                            st.caption(f"{below_count} below 3.0")
                        else:
                            st.caption("All ≥ 3.0")
        else:
            st.info("No field evaluations with disposition data available")
    else:
        st.info("No field evaluations available for disposition analysis")
    
    # Enhanced Recent Evaluations Table
    st.subheader("📋 Recent Evaluations")
    
    # Check if created_at column exists and sort accordingly
    if 'created_at' in df.columns:
        recent_df = df.sort_values('created_at', ascending=False).head(10)
    else:
        recent_df = df.head(10)
    
    # Enhanced display columns with new dashboard fields
    display_columns = ['student_name', 'evaluator_name', 'school_name', 'subject_area', 
                      'department', 'semester', 'rubric_type', 'status', 'total_score']
    
    # Only include columns that exist
    available_columns = [col for col in display_columns if col in recent_df.columns]
    display_df = recent_df[available_columns].copy()
    
    # Add date if available
    if 'created_at' in df.columns:
        display_df['created_at'] = pd.to_datetime(recent_df['created_at']).dt.strftime('%Y-%m-%d')
        available_columns.append('created_at')
    
    # Enhanced column configuration with new dashboard fields
    column_config = {
        "student_name": st.column_config.TextColumn("Student", width="medium"),
        "evaluator_name": st.column_config.TextColumn("Evaluator", width="medium"),
        "school_name": st.column_config.TextColumn("School", width="large"),
        "subject_area": st.column_config.TextColumn("Subject", width="medium"),
        "department": st.column_config.TextColumn("Department", width="small"),
        "semester": st.column_config.TextColumn("Semester", width="small"),
        "rubric_type": st.column_config.TextColumn("Type", width="medium"),
        "status": st.column_config.TextColumn("Status", width="small"),
        "total_score": st.column_config.NumberColumn("Score", format="%d", width="small"),
        "created_at": st.column_config.DateColumn("Date", width="medium")
    }
    
    # Filter column config to only available columns
    filtered_config = {k: v for k, v in column_config.items() if k in available_columns}
    
    st.dataframe(
        display_df,
        column_config=filtered_config,
        hide_index=True,
        use_container_width=True
    )
    
    # Detailed Evaluation Viewer
    st.subheader("🔍 Detailed Evaluation Viewer")
    
    if completed_evals:
        eval_options = [f"{e['student_name']} - {e.get('school_name', 'Unknown School')} ({e['created_at'][:10]})" 
                       for e in completed_evals]
        
        selected_eval_idx = st.selectbox("Select evaluation to view details:", 
                                        range(len(eval_options)), 
                                        format_func=lambda x: eval_options[x])
        
        if selected_eval_idx is not None:
            selected_eval = completed_evals[selected_eval_idx]
            show_detailed_evaluation_view(selected_eval)
    
    # Export functionality
    st.subheader("💾 Data Management")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📥 Export All Data"):
            export_data = {
                'evaluations': evaluations,
                'export_date': datetime.now().isoformat(),
                'version': '1.0',
                'summary': {
                    'total_evaluations': len(evaluations),
                    'completed_evaluations': completed,
                    'average_score': avg_score if completed > 0 else 0
                }
            }
            st.download_button(
                "Download JSON",
                json.dumps(export_data, indent=2),
                f"ster-evaluations-{datetime.now().strftime('%Y%m%d')}.json",
                "application/json"
            )
    
    with col2:
        uploaded_file = st.file_uploader("📤 Import Data", type=['json'])
        if uploaded_file:
            try:
                data = json.load(uploaded_file)
                imported_count = import_data(data)
                st.success(f"Imported {imported_count} evaluations!")
                st.rerun()
            except Exception as e:
                st.error(f"Import failed: {str(e)}")

def analyze_competency_performance(evaluations):
    """Analyze performance by competency area"""
    from data.rubrics import get_field_evaluation_items, get_ster_items, filter_items_by_evaluator_role
    
    competency_scores = {}
    competency_counts = {}
    
    for evaluation in evaluations:
        rubric_type = evaluation.get('rubric_type', 'field_evaluation')
        scores = evaluation.get('scores', {})
        
        # Get appropriate rubric items
        if rubric_type == 'field_evaluation':
            items = get_field_evaluation_items()
        else:
            # Get STER items filtered for supervisors
            all_ster_items = get_ster_items()
            items = filter_items_by_evaluator_role(all_ster_items, 'supervisor')
        
        for item in items:
            item_id = item['id']
            competency_area = item['competency_area']
            
            if item_id in scores:
                if competency_area not in competency_scores:
                    competency_scores[competency_area] = 0
                    competency_counts[competency_area] = 0
                
                competency_scores[competency_area] += scores[item_id]
                competency_counts[competency_area] += 1
    
    # Calculate averages
    competency_averages = {}
    for area in competency_scores:
        if competency_counts[area] > 0:
            competency_averages[area] = competency_scores[area] / competency_counts[area]
    
    return competency_averages

def analyze_disposition_performance(evaluations):
    """Analyze professional disposition performance"""
    from data.rubrics import get_professional_dispositions
    
    dispositions = get_professional_dispositions()
    disposition_scores = {}
    disposition_counts = {}
    
    for evaluation in evaluations:
        disp_scores = evaluation.get('disposition_scores', {})
        
        for disposition in dispositions:
            disp_id = disposition['id']
            disp_name = disposition['name']
            
            if disp_id in disp_scores:
                if disp_name not in disposition_scores:
                    disposition_scores[disp_name] = 0
                    disposition_counts[disp_name] = 0
                
                disposition_scores[disp_name] += disp_scores[disp_id]
                disposition_counts[disp_name] += 1
    
    # Calculate averages
    disposition_averages = {}
    for disp_name in disposition_scores:
        if disposition_counts[disp_name] > 0:
            disposition_averages[disp_name] = disposition_scores[disp_name] / disposition_counts[disp_name]
    
    return disposition_averages

def show_detailed_evaluation_view(evaluation):
    """Show detailed view of a specific evaluation"""
    from data.rubrics import get_field_evaluation_items, get_ster_items, get_professional_dispositions, filter_items_by_evaluator_role
    
    # Basic information
    col1, col2, col3 = st.columns(3)
    with col1:
        st.info(f"**Student:** {evaluation['student_name']}")
        st.info(f"**Evaluator:** {evaluation['evaluator_name']}")
    with col2:
        st.info(f"**School:** {evaluation.get('school_name', 'N/A')}")
        st.info(f"**Subject:** {evaluation.get('subject_area', 'N/A')}")
    with col3:
        st.info(f"**Type:** {evaluation['rubric_type'].replace('_', ' ').title()}")
        st.info(f"**Total Score:** {evaluation['total_score']}")
    
    # Assessment Items Detail
    rubric_type = evaluation['rubric_type']
    if rubric_type == 'field_evaluation':
        items = get_field_evaluation_items()
    else:
        # Get STER items filtered for supervisors
        all_ster_items = get_ster_items()
        items = filter_items_by_evaluator_role(all_ster_items, 'supervisor')
    scores = evaluation.get('scores', {})
    justifications = evaluation.get('justifications', {})
    
    st.subheader("Assessment Items Breakdown")
    
    for item in items:
        item_id = item['id']
        score = scores.get(item_id, 0)
        justification = justifications.get(item_id, 'No justification provided')
        
        with st.expander(f"{item['code']}: {item['title']} (Score: {score})"):
            col1, col2 = st.columns([1, 2])
            with col1:
                st.write(f"**Competency Area:** {item['competency_area']}")
                st.write(f"**Score:** Level {score}")
                
                # Color code the score
                if score >= 3:
                    st.success("Exceeds expectations")
                elif score >= 2:
                    st.info("Meets expectations")
                elif score >= 1:
                    st.warning("Approaching expectations")
                else:
                    st.error("Does not demonstrate")
            
            with col2:
                st.write("**Justification:**")
                st.write(justification)
    
    # Dispositions Detail
    st.subheader("Professional Dispositions")
    
    dispositions = get_professional_dispositions()
    disposition_scores = evaluation.get('disposition_scores', {})
    disposition_comments = evaluation.get('disposition_comments', {})
    
    for disposition in dispositions:
        disp_id = disposition['id']
        score = disposition_scores.get(disp_id, 0)
        comment = disposition_comments.get(disp_id, "No feedback provided")
        
        col1, col2 = st.columns([1, 2])
        with col1:
            st.write(f"**{disposition['name']}**")
            st.caption(disposition['description'])
            if score >= 3:
                st.success(f"Level {score}")
            else:
                st.error(f"Level {score} (Needs Level 3+)")
        with col2:
            st.write("**Supervisor Feedback:**")
            st.write(comment if comment.strip() else "_No feedback provided_")
    
    # PDF Export button for this evaluation
    st.markdown("---")
    st.subheader("📄 Export Options")
    col1_export, col2_export = st.columns(2)
    
    with col1_export:
        # Prepare data for PDF generation
        pdf_data = {
            'rubric_type': evaluation['rubric_type'],
            'student_name': evaluation['student_name'],
            'evaluator_name': evaluation['evaluator_name'],
            'date': evaluation.get('created_at', datetime.now().isoformat())[:10],
            'school': evaluation.get('school_name', 'N/A'),
            'subject': evaluation.get('subject_area', 'N/A'),
            'is_formative': True,  # TODO: Add formative/summative tracking
            'competency_scores': [],
            'total_items': len(items),
            'meeting_expectations': sum(1 for score in scores.values() if isinstance(score, int) and score >= 2),
            'areas_for_growth': sum(1 for score in scores.values() if isinstance(score, int) and score < 2)
        }
        
        # Add competency scores
        for item in items:
            score = scores.get(item['id'], 0)
            pdf_data['competency_scores'].append({
                'competency': f"{item['code']}: {item['title']}",
                'score': score,
                'justification': justifications.get(item['id'], 'No justification provided')
            })
        
        # Add dispositions
        if evaluation['rubric_type'] == 'field_evaluation' and disposition_scores:
            pdf_data['dispositions'] = []
            for disposition in dispositions:
                disp_id = disposition['id']
                score = disposition_scores.get(disp_id, 0)
                pdf_data['dispositions'].append({
                    'disposition': disposition['name'],
                    'score': score,
                    'notes': disposition_comments.get(disp_id, '')
                })
        
        # Add AI analysis if available
        ai_analyses = evaluation.get('ai_analyses', {})
        if ai_analyses:
            # Collect all AI analyses
            all_analyses = []
            strengths = []
            areas_for_growth = []
            
            for item_id, analysis in ai_analyses.items():
                if analysis and isinstance(analysis, str):
                    # Get the item name for context
                    item = next((i for i in items if i['id'] == item_id), None)
                    if item:
                        item_name = f"{item['code']}: {item['title']}"
                        # Add the analysis with context
                        all_analyses.append(f"{item_name}\n{analysis}")
                        
                        # Extract strengths (look for positive indicators)
                        if any(word in analysis.lower() for word in ['strong', 'excellent', 'effective', 'well', 'good', 'demonstrates']):
                            strengths.append(f"{item['code']}: {analysis[:150]}...")
                        
                        # Extract areas for growth (look for improvement indicators)
                        if any(word in analysis.lower() for word in ['improve', 'develop', 'consider', 'could', 'should', 'needs']):
                            areas_for_growth.append(f"{item['code']}: {analysis[:150]}...")
            
            pdf_data['ai_analysis'] = {
                'strengths': strengths[:5],  # Limit to top 5
                'areas_for_growth': areas_for_growth[:5],  # Limit to top 5
                'recommendations': [],  # Could be extracted from overall analysis
                'full_analyses': all_analyses  # Include all analyses
            }
        
        # Generate PDF
        try:
            pdf_bytes = pdf_service.generate_evaluation_pdf(pdf_data)
            
            # Create filename
            filename = f"{evaluation['student_name'].replace(' ', '_')}_{evaluation['rubric_type']}_{evaluation.get('created_at', datetime.now().isoformat())[:10]}.pdf"
            
            st.download_button(
                label="📄 Download Evaluation Report (PDF)",
                data=pdf_bytes,
                file_name=filename,
                mime="application/pdf",
                help="Download the complete evaluation report as a PDF file"
            )
        except Exception as e:
            st.error(f"Error generating PDF: {str(e)}")
    
    with col2_export:
        st.info("💡 **Tip**: Download the PDF report for archival or distribution purposes.")
    
    # AI Performance Evaluation Comparison View
    if st.session_state.get('show_ai_comparison', False):
        st.markdown("---")
        st.subheader("🤖 AI Performance Evaluation - Comparison Report")
        st.caption("Showing AI-generated content vs. Supervisor revisions")
        
        # Close button
        if st.button("✖️ Close Comparison", key="close_comparison"):
            st.session_state.show_ai_comparison = False
            st.rerun()
        
        if st.session_state.get('ai_original_data'):
            ai_data = st.session_state.ai_original_data
            current_data = {
                'justifications': st.session_state.get('justifications', {}),
                'scores': st.session_state.get('scores', {}),
                'observation_notes': st.session_state.get('observation_notes', '')
            }
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                ai_saved_time = ai_data.get('saved_at', 'Unknown')
                st.metric("AI Version Saved", ai_saved_time[:19] if len(ai_saved_time) > 19 else ai_saved_time)
            with col2:
                total_items = len(items)
                st.metric("Total Competencies", total_items)
            with col3:
                modified_count = sum(1 for item_id in ai_data.get('justifications', {})
                                   if ai_data['justifications'].get(item_id, '') != current_data['justifications'].get(item_id, ''))
                st.metric("Modified Justifications", modified_count)
            with col4:
                score_changes = sum(1 for item_id in ai_data.get('scores', {})
                                  if ai_data['scores'].get(item_id) != current_data['scores'].get(item_id))
                st.metric("Score Changes", score_changes)
            
            # Detailed comparison for each competency
            st.markdown("### 📋 Competency-by-Competency Comparison")
            
            for item in items:
                item_id = item['id']
                ai_justification = ai_data.get('justifications', {}).get(item_id, '')
                current_justification = current_data['justifications'].get(item_id, '')
                ai_score = ai_data.get('scores', {}).get(item_id)
                current_score = current_data['scores'].get(item_id)
                
                # Check if there are differences
                has_changes = (ai_justification != current_justification) or (ai_score != current_score)
                
                with st.expander(f"{item['name']} {'🔄' if has_changes else '✓'}", expanded=has_changes):
                    # Score comparison
                    if ai_score is not None or current_score is not None:
                        score_col1, score_col2 = st.columns(2)
                        with score_col1:
                            st.markdown("**AI Score:**")
                            if ai_score is not None:
                                st.write(f"Level {ai_score}")
                            else:
                                st.write("Not scored")
                        with score_col2:
                            st.markdown("**Final Score:**")
                            if current_score is not None:
                                st.write(f"Level {current_score}")
                                if ai_score is not None and ai_score != current_score:
                                    change = current_score - ai_score
                                    st.caption(f"{'↑' if change > 0 else '↓'} {abs(change)} level{'s' if abs(change) > 1 else ''}")
                            else:
                                st.write("Not scored")
                    
                    # Justification comparison
                    st.markdown("---")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**🤖 AI-Generated Justification:**")
                        if ai_justification:
                            st.text_area(f"AI: {item_id}", ai_justification, height=200, disabled=True, key=f"ai_just_{item_id}")
                        else:
                            st.info("No AI justification generated")
                    
                    with col2:
                        st.markdown("**👨‍🏫 Supervisor Final Version:**")
                        if current_justification:
                            st.text_area(f"Final: {item_id}", current_justification, height=200, disabled=True, key=f"final_just_{item_id}")
                        else:
                            st.info("No justification provided")
                    
                    # Highlight differences
                    if has_changes and ai_justification and current_justification:
                        st.caption("💡 Supervisor made modifications to this competency")
            
            # Export comparison data
            st.markdown("### 📊 Export Comparison Data")
            comparison_data = {
                'evaluation_info': {
                    'student_name': student_name,
                    'evaluator_name': evaluator_name,
                    'evaluation_date': evaluation_date.isoformat() if 'evaluation_date' in locals() else datetime.now().isoformat(),
                    'rubric_type': rubric_type
                },
                'ai_original': ai_data,
                'supervisor_final': current_data,
                'comparison_summary': {
                    'total_competencies': len(items),
                    'modified_justifications': modified_count,
                    'score_changes': score_changes,
                    'ai_version_saved_at': ai_data.get('saved_at', 'Unknown')
                },
                'export_timestamp': datetime.now().isoformat()
            }
            
            comparison_json = json.dumps(comparison_data, indent=2)
            
            st.download_button(
                "📥 Download Comparison Data (JSON)",
                comparison_json,
                f"ai_comparison_{student_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                "application/json",
                help="Download the complete comparison data for research analysis"
            )

def show_research_comparison():
    """Research comparison view for AI vs Supervisor modifications"""
    st.header("🔬 Research Comparison: AI vs Supervisor Modifications")
    
    # Get all evaluations with AI originals
    evaluations = load_evaluations()
    ai_evaluations = [e for e in evaluations if e.get('has_ai_original', False)]
    
    if not ai_evaluations:
        st.info("No evaluations with saved AI versions found.")
        st.write("To use this feature:")
        st.write("1. Create a new evaluation")
        st.write("2. Generate AI analysis")
        st.write("3. Make your modifications to the evaluation")
        st.write("4. Complete the evaluation")
        return
    
    # Selection dropdown
    selected_eval = st.selectbox(
        "Select an evaluation to compare:",
        ai_evaluations,
        format_func=lambda e: f"{e['student_name']} - {e['date']} - {e.get('rubric_type', 'Unknown').replace('_', ' ').title()}"
    )
    
    if selected_eval:
        comparison = get_evaluation_comparison(selected_eval['id'])
        
        if comparison and comparison['has_changes']:
            st.success(f"Found {len(comparison['differences'])} differences between AI and supervisor versions")
            
            # Display comparison metrics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total Changes", len(comparison['differences']))
            with col2:
                score_changes = len([d for d in comparison['differences'] if d['field'] == 'score'])
                st.metric("Score Changes", score_changes)
            with col3:
                just_changes = len([d for d in comparison['differences'] if d['field'] == 'justification'])
                st.metric("Justification Changes", just_changes)
            
            # Detailed comparison
            st.markdown("### 📋 Detailed Comparison")
            
            # Group by competency
            from data.rubrics import get_field_evaluation_items, get_ster_items
            items = get_field_evaluation_items() if selected_eval['rubric_type'] == 'field_evaluation' else get_ster_items()
            item_dict = {item['id']: item for item in items}
            
            for diff in comparison['differences']:
                item_id = diff['item_id']
                item = item_dict.get(item_id, {})
                
                with st.expander(f"{item.get('name', 'Unknown')} - {diff['field'].title()} Changed"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**🤖 AI Generated:**")
                        if diff['field'] == 'score':
                            st.write(f"Score: {diff['ai_value']}")
                        else:
                            st.text_area("AI Justification", diff['ai_value'], height=150, disabled=True)
                    
                    with col2:
                        st.markdown("**👨‍🏫 Supervisor Modified:**")
                        if diff['field'] == 'score':
                            st.write(f"Score: {diff['current_value']}")
                            if diff['ai_value'] != diff['current_value']:
                                change = diff['current_value'] - diff['ai_value']
                                st.write(f"Change: {'↑' if change > 0 else '↓'} {abs(change)}")
                        else:
                            st.text_area("Supervisor Justification", diff['current_value'], height=150, disabled=True)
            
            # Export comparison data
            st.markdown("### 📊 Export Data")
            export_data = {
                'evaluation': selected_eval,
                'comparison': comparison,
                'export_date': datetime.now().isoformat()
            }
            
            st.download_button(
                "Download Comparison JSON",
                json.dumps(export_data, indent=2),
                f"ai_comparison_{selected_eval['student_name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.json",
                "application/json"
            )
            
        else:
            st.info("No differences found between AI and current versions.")
            st.write("This could mean:")
            st.write("- The supervisor accepted all AI-generated content")
            st.write("- The evaluation hasn't been completed yet")

def show_evaluation_form():
    """Evaluation form for creating new evaluations"""
    st.header("📝 New Evaluation")
    
    # Get dispositions (same for all evaluation types)
    dispositions = get_professional_dispositions()
    
    # Initialize session state for lesson plan analysis
    if 'lesson_plan_analysis' not in st.session_state:
        st.session_state.lesson_plan_analysis = None
    if 'extracted_info' not in st.session_state:
        st.session_state.extracted_info = {}
    
    # STEP 1: Lesson Plan Upload (Optional)
    st.subheader("📄 Step 1: Upload Lesson Plan (Optional)")
    st.caption("Upload the student teacher's lesson plan to automatically extract evaluation information, or skip to proceed without it")
    
    # Add skip option
    col1, col2 = st.columns([3, 1])
    with col1:
        # Option to use synthetic data or upload real file
        input_method = st.radio(
            "Choose input method:",
            ["📤 Upload File", "🧪 Use Synthetic Data", "✏️ Paste Text", "⏭️ Skip Lesson Plan"],
            horizontal=True
        )
    
    with col2:
        if input_method != "⏭️ Skip Lesson Plan":
            st.info("💡 Lesson plan helps AI generate better analysis")
        else:
            st.warning("⚠️ Proceeding without lesson plan")
    
    lesson_plan_text = None
    
    if input_method == "⏭️ Skip Lesson Plan":
        st.info("✅ Skipping lesson plan upload. You can proceed directly to basic information.")
        lesson_plan_text = None
        # Clear any existing lesson plan analysis
        st.session_state.lesson_plan_analysis = None
        
    elif input_method == "🧪 Use Synthetic Data":
        # Load available synthetic lesson plans
        evaluations = load_evaluations()
        synthetic_evaluations = [e for e in evaluations if e.get('is_synthetic', False) and e.get('lesson_plan')]
        
        if synthetic_evaluations:
            st.info(f"Found {len(synthetic_evaluations)} synthetic lesson plans available for testing")
            
            # Create selection options
            eval_options = []
            for i, eval_data in enumerate(synthetic_evaluations):
                student_name = eval_data.get('student_name', f'Student {i+1}')
                subject = eval_data.get('subject_area', 'Unknown Subject')
                grade = eval_data.get('grade_levels', 'Unknown Grade')
                school = eval_data.get('school_name', 'Unknown School')
                option_text = f"{student_name} - {subject} ({grade}) at {school}"
                eval_options.append(option_text)
            
            selected_index = st.selectbox(
                "Select a synthetic lesson plan:",
                range(len(eval_options)),
                format_func=lambda x: eval_options[x],
                help="Choose from available synthetic lesson plans for testing"
            )
            
            if selected_index is not None:
                selected_evaluation = synthetic_evaluations[selected_index]
                lesson_plan_text = selected_evaluation.get('lesson_plan', '')
                
                # Show preview
                with st.expander("📋 Preview Selected Lesson Plan"):
                    st.text_area(
                        "Lesson Plan Content:",
                        value=lesson_plan_text[:500] + "..." if len(lesson_plan_text) > 500 else lesson_plan_text,
                        height=200,
                        disabled=True
                    )
                
                st.success(f"✅ Loaded synthetic lesson plan for {selected_evaluation.get('student_name', 'Unknown Student')}")
        else:
            st.warning("No synthetic lesson plans available. Generate some test data first!")
            if st.button("🧪 Go to Test Data Generation"):
                st.session_state.page = "🧪 Test Data"
                st.rerun()
    
    elif input_method == "📤 Upload File":
        uploaded_file = st.file_uploader(
            "Choose lesson plan file",
            type=['txt', 'docx', 'pdf', 'doc'],
            help="Upload the lesson plan in text, Word, or PDF format"
        )
    
        if uploaded_file is not None:
            # Read file content based on type with automatic processing
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            try:
                if uploaded_file.type == "text/plain" or file_extension == 'txt':
                    # Handle text files
                    lesson_plan_text = str(uploaded_file.read(), "utf-8")
                    st.success(f"✅ Text file '{uploaded_file.name}' processed successfully!")
                    
                elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                                          "application/msword"] or file_extension == 'docx':
                    # Handle DOCX files with automatic text extraction
                    try:
                        from docx import Document
                        
                        # Extract text from Word document with enhanced processing
                        doc = Document(uploaded_file)
                        lesson_plan_text = ""
                        
                        # Extract text from paragraphs
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                lesson_plan_text += paragraph.text + "\n"
                        
                        # Extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        lesson_plan_text += cell.text + " "
                                lesson_plan_text += "\n"
                        
                        # Extract text from headers and footers
                        for section in doc.sections:
                            if hasattr(section, 'header') and section.header:
                                for paragraph in section.header.paragraphs:
                                    if paragraph.text.strip():
                                        lesson_plan_text += paragraph.text + "\n"
                            if hasattr(section, 'footer') and section.footer:
                                for paragraph in section.footer.paragraphs:
                                    if paragraph.text.strip():
                                        lesson_plan_text += paragraph.text + "\n"
                        
                        if lesson_plan_text.strip():
                            st.success(f"✅ Word document '{uploaded_file.name}' processed successfully!")
                            # Show preview
                            with st.expander("📄 Preview extracted content"):
                                st.text_area(
                                    "Extracted Lesson Plan Content",
                                    value=lesson_plan_text,
                                    height=150,
                                    disabled=True,
                                    key="lesson_plan_docx_preview"
                                )
                        else:
                            st.warning("⚠️ No text could be extracted from this Word document.")
                            lesson_plan_text = st.text_area(
                                "Paste lesson plan content:",
                                height=200,
                                placeholder="Copy and paste your lesson plan content here..."
                            )
                            
                    except ImportError:
                        st.error("❌ Word document processing library not available.")
                        lesson_plan_text = st.text_area(
                            "Paste lesson plan content:",
                            height=200,
                            placeholder="Copy and paste your lesson plan content here..."
                        )
                    except Exception as e:
                        st.warning(f"⚠️ Standard processing failed: {str(e)}")
                        
                        # Try alternative extraction method for corrupted files
                        st.info("🔧 Trying alternative extraction method for corrupted/problematic files...")
                        try:
                            import docx2txt
                            uploaded_file.seek(0)  # Reset file pointer
                            lesson_plan_text = docx2txt.process(uploaded_file)
                            
                            if lesson_plan_text and lesson_plan_text.strip():
                                st.success(f"✅ Alternative method successfully extracted text from '{uploaded_file.name}'!")
                                # Show preview
                                with st.expander("📄 Preview extracted content"):
                                    st.text_area(
                                        "Extracted Lesson Plan Content",
                                        value=lesson_plan_text,
                                        height=150,
                                        disabled=True,
                                        key="lesson_plan_docx2txt_preview"
                                    )
                            else:
                                raise Exception("No text could be extracted")
                                
                        except Exception as e2:
                            st.error(f"❌ Both extraction methods failed: {str(e2)}")
                            st.info("💡 **Troubleshooting tips for corrupted files:**")
                            st.markdown("""
                            - **Font Issues**: Try opening in Word and saving as new .docx file
                            - **Google Docs Method**: Upload to Google Drive → Open with Google Docs → Copy text
                            - **Online Converters**: Try SmallPDF, Zamzar, or CloudConvert
                            - **Document Repair**: In Word, use File → Open → Open and Repair
                            - **Manual Option**: Use the text area below as fallback
                            """)
                            lesson_plan_text = st.text_area(
                                "Paste lesson plan content:",
                                height=200,
                                placeholder="Copy and paste your lesson plan content here..."
                            )
                        
                elif uploaded_file.type == "application/pdf" or file_extension == 'pdf':
                    # Handle PDF files with automatic text extraction
                    try:
                        import pdfplumber
                        
                        # Extract text using pdfplumber
                        with pdfplumber.open(uploaded_file) as pdf:
                            lesson_plan_text = ""
                            for page in pdf.pages:
                                text = page.extract_text()
                                if text:
                                    lesson_plan_text += text + "\n"
                        
                        if lesson_plan_text.strip():
                            st.success(f"✅ PDF '{uploaded_file.name}' processed successfully!")
                            # Show preview
                            with st.expander("📄 Preview extracted content"):
                                st.text_area(
                                    "Extracted Lesson Plan Content",
                                    value=lesson_plan_text,
                                    height=150,
                                    disabled=True,
                                    key="lesson_plan_pdf_preview"
                                )
                        else:
                            st.warning("⚠️ No text could be extracted from this PDF.")
                            lesson_plan_text = st.text_area(
                                "Paste lesson plan content:",
                                height=200,
                                placeholder="Copy and paste your lesson plan content here..."
                            )
                            
                    except ImportError:
                        st.error("❌ PDF processing libraries not available.")
                        lesson_plan_text = st.text_area(
                            "Paste lesson plan content:",
                            height=200,
                            placeholder="Copy and paste your lesson plan content here..."
                        )
                    except Exception as e:
                        st.error(f"❌ Error processing PDF: {str(e)}")
                        lesson_plan_text = st.text_area(
                            "Paste lesson plan content:",
                            height=200,
                            placeholder="Copy and paste your lesson plan content here..."
                        )
                        
                else:
                    # Unknown file type - fallback to manual entry
                    st.warning(f"⚠️ File type not recognized. Please paste the content manually.")
                    lesson_plan_text = st.text_area(
                        "Paste lesson plan content:",
                        height=200,
                        placeholder="Copy and paste your lesson plan content here..."
                    )
                    
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
                lesson_plan_text = st.text_area(
                    "Paste lesson plan content:",
                    height=200,
                    placeholder="Copy and paste your lesson plan content here..."
                )
    
    elif input_method == "✏️ Paste Text":
        lesson_plan_text = st.text_area(
            "Paste lesson plan content directly:",
            height=200,
            placeholder="Paste your lesson plan content here...",
            help="Copy and paste the lesson plan text for AI analysis"
        )
    
    # AI Analysis of Lesson Plan
    if lesson_plan_text and len(lesson_plan_text.strip()) > 50:
        if openai_service.is_enabled():
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("🤖 Analyze Lesson Plan with AI", type="primary"):
                    with st.spinner("Analyzing lesson plan..."):
                        try:
                            analysis = openai_service.analyze_lesson_plan(lesson_plan_text)
                            analysis['extraction_timestamp'] = datetime.now().isoformat()
                            st.session_state.lesson_plan_analysis = analysis
                            
                            # Check confidence score
                            confidence = analysis.get('confidence_score', 0)
                            if confidence == 0:
                                st.warning("⚠️ AI analysis completed but no information could be extracted. Please check the lesson plan format.")
                            elif confidence < 0.3:
                                st.warning(f"⚠️ Limited information extracted (confidence: {confidence:.1%}). Some fields may need manual entry.")
                            else:
                                st.success("✅ Lesson plan analyzed successfully!")
                            
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ AI Analysis Error: {str(e)}")
                            st.info("💡 **Common issues:**")
                            st.markdown("""
                            - Check that your OpenAI API key is valid
                            - Ensure the lesson plan contains readable text
                            - Try a shorter or simpler document format
                            - Check Settings to verify AI is properly configured
                            """)
            
            with col2:
                if st.session_state.lesson_plan_analysis:
                    st.success("✅ Analysis Complete")
                    confidence = st.session_state.lesson_plan_analysis.get('confidence_score', 0)
                    st.metric("AI Confidence", f"{confidence:.1%}")
        else:
            st.warning("🤖 AI features disabled. Add OpenAI API key in Settings to enable automatic extraction.")
    
    # Display Extracted Information
    if st.session_state.lesson_plan_analysis:
        st.subheader("📋 Step 2: Review Extracted Information")
        st.caption("Review and modify the information extracted from the lesson plan. Supervisor notes override lesson plan data.")
        
        analysis = st.session_state.lesson_plan_analysis
        
        # Create two columns for extracted vs editable info
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.markdown("**🤖 AI Extracted Information**")
            with st.container():
                st.text_input("Teacher Name (Extracted)", value=analysis.get('teacher_name', 'Not found'), disabled=True)
                st.text_input("Lesson Date (Extracted)", value=analysis.get('lesson_date', 'Not found'), disabled=True)
                st.text_input("Subject (Extracted)", value=analysis.get('subject_area', 'Not found'), disabled=True)
                st.text_input("Grade Level (Extracted)", value=analysis.get('grade_levels', 'Not found'), disabled=True)
                st.text_input("School (Extracted)", value=analysis.get('school_name', 'Not found'), disabled=True)
                st.text_input("Class Size (Extracted)", value=str(analysis.get('total_students', 'Not found')), disabled=True)
        
        with col2:
            st.markdown("**✏️ Supervisor Override (Fill as needed)**")
            with st.container():
                student_name = st.text_input(
                    "Student Teacher Name *", 
                    value=analysis.get('teacher_name', ''),
                    key="student_name",
                    help="Override the extracted teacher name if incorrect"
                )
                
                evaluation_date = st.date_input(
                    "Evaluation Date *",
                    value=datetime.now().date(),
                    help="Date of this evaluation (may differ from lesson date)"
                )
                
                subject_area = st.text_input(
                    "Subject Area",
                    value=analysis.get('subject_area', ''),
                    key="subject_area_override"
                )
                
                department = st.selectbox(
                    "Department *", 
                    ["Elementary", "Secondary", "GRAD", "Special Ed"], 
                    index=1, key="department_override"
                )
                
                semester = st.selectbox(
                    "Current Semester *", 
                    ["Fall 2024", "Spring 2025", "Summer 2025", "Fall 2025"], 
                    index=1, key="semester_override"
                )
                
                grade_levels = st.text_input(
                    "Grade Levels",
                    value=analysis.get('grade_levels', ''),
                    key="grade_levels_override",
                    help="Optional: Specific grade levels"
                )
                
                school_name = st.text_input(
                    "School Name",
                    value=analysis.get('school_name', ''),
                    key="school_name_override"
                )
                
                class_size = st.number_input(
                    "Class Size",
                    min_value=1,
                    max_value=40,
                    value=analysis.get('total_students', 20) if analysis.get('total_students') else 20,
                    key="class_size_override"
                )
                
                # Additional context from lesson plan
                lesson_topic = st.text_input(
                    "Lesson Topic",
                    value=analysis.get('lesson_topic', ''),
                    key="lesson_topic_override"
                )
        
        # Store the final information (supervisor override takes precedence)
        st.session_state.extracted_info = {
            'student_name': student_name,
            'subject_area': subject_area or analysis.get('subject_area', ''),
            'department': department,
            'semester': semester,
            'grade_levels': grade_levels or analysis.get('grade_levels', ''),
            'school_name': school_name or analysis.get('school_name', ''),
            'class_size': class_size,
            'lesson_topic': lesson_topic or analysis.get('lesson_topic', ''),
            'evaluation_date': evaluation_date.isoformat(),
            'lesson_plan_text': lesson_plan_text,
            'ai_analysis': analysis
        }
    
    # Basic Information (Traditional Form - fallback if no lesson plan)
    else:
        st.subheader("📋 Step 2: Basic Information")
        if input_method == "⏭️ Skip Lesson Plan":
            st.caption("✅ No lesson plan provided - please enter evaluation information manually")
            # Visual indicator for skipped lesson plan
            st.info("📋 **Lesson Plan Status:** Skipped - AI analysis will rely solely on observation notes")
        else:
            st.caption("Enter evaluation information manually")
        
        # Add synthetic data button
        if st.button("🧪 Fill with Test Data", key="fill_synthetic_data", 
                    help="Auto-fill form with synthetic data for testing"):
            import random
            from data.sample_observation_notes import get_sample_observation_notes
            
            # Generate synthetic names
            student_first_names = ["Sarah", "Michael", "Jessica", "David", "Emily", "John", "Amanda", "Robert"]
            student_last_names = ["Johnson", "Smith", "Williams", "Brown", "Davis", "Miller", "Wilson", "Moore"]
            
            supervisor_first_names = ["Dr. Patricia", "Dr. James", "Dr. Linda", "Dr. William", "Dr. Susan", "Dr. Richard"]
            supervisor_last_names = ["Anderson", "Thompson", "Martinez", "Taylor", "Clark", "Rodriguez", "Lewis", "Walker"]
            
            # Set synthetic data in session state
            st.session_state.student_name = f"{random.choice(student_first_names)} {random.choice(student_last_names)}"
            st.session_state.evaluator_name = f"{random.choice(supervisor_first_names)} {random.choice(supervisor_last_names)}"
            st.session_state.subject_area_manual = random.choice(["Mathematics", "Science", "English Language Arts", "Social Studies"])
            st.session_state.school_name_manual = random.choice(["Lincoln Elementary", "Washington Middle School", "Jefferson High School", "Roosevelt Academy"])
            st.session_state.grade_levels_manual = random.choice(["3rd Grade", "4th Grade", "5th Grade", "6th Grade", "7th Grade", "8th Grade"])
            st.session_state.class_size_manual = random.randint(18, 28)
            st.session_state.observation_notes = get_sample_observation_notes()
            
            st.success("✅ Test data filled! Observation notes have been pre-populated.")
            st.rerun()
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            student_name = st.text_input("Student Teacher Name *", key="student_name")
        with col2:
            evaluation_date = st.date_input("Evaluation Date *", value=datetime.now().date())
        with col3:
            subject_area = st.text_input("Subject Area", key="subject_area_manual")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            department = st.selectbox("Department *", 
                                    ["Elementary", "Secondary", "GRAD", "Special Ed"], 
                                    index=1, key="department_manual")  # Default to Secondary
        with col2:
            school_name = st.text_input("School Name", key="school_name_manual")
        with col3:
            semester = st.selectbox("Current Semester *", 
                                  ["Fall 2024", "Spring 2025", "Summer 2025", "Fall 2025"], 
                                  index=1, key="semester_manual")  # Default to Spring 2025
        
        col1, col2 = st.columns(2)
        with col1:
            grade_levels = st.text_input("Grade Level(s)", key="grade_levels_manual", 
                                       placeholder="e.g., 3rd-5th, 9th-12th",
                                       help="Optional: Specific grade levels")
        with col2:
            class_size = st.number_input("Class Size", min_value=1, max_value=40, value=20)
        
        # Visual reminder about lesson plan benefits
        if input_method == "⏭️ Skip Lesson Plan":
            with st.expander("💡 Why upload a lesson plan?", expanded=False):
                st.markdown("""
                **Benefits of uploading a lesson plan:**
                - 🤖 **Better AI Analysis**: AI can compare planned vs. observed activities
                - 📊 **More Accurate Justifications**: Evidence-based analysis using lesson objectives
                - ⚡ **Faster Evaluation**: Auto-extraction of student and lesson information
                - 🎯 **Contextual Insights**: AI understands the lesson's goals and structure
                
                **You can still get good results without a lesson plan** - just ensure your observation notes are detailed!
                """)
        
        # Store manual information
        st.session_state.extracted_info = {
            'student_name': student_name,
            'subject_area': subject_area,
            'department': department,
            'semester': semester,
            'grade_levels': grade_levels,
            'school_name': school_name,
            'class_size': class_size,
            'evaluation_date': evaluation_date.isoformat(),
            'lesson_plan_text': lesson_plan_text if lesson_plan_text else None,
            'ai_analysis': None
        }
    
    # Evaluator Information
    st.subheader("👨‍🏫 Evaluator Information")
    col1, col2 = st.columns(2)
    with col1:
        evaluator_name = st.text_input("Evaluator Name *", key="evaluator_name")
    with col2:
        # Supervisor is the only role now
        evaluator_role = "supervisor"
        st.text_input("Evaluator Role", value="Supervisor", disabled=True)
    
    # Evaluation Type Selection
    st.subheader("📋 Evaluation Type")
    rubric_type = st.selectbox(
        "Select Evaluation Type",
        ["field_evaluation", "ster"],
        index=1,  # Make STER default (index 1)
        format_func=lambda x: "Field Evaluation" if x == "field_evaluation" else "STER"
    )
    
    # Add rubric reference section
    with st.expander("📖 **View Official Rubric**", expanded=False):
        if rubric_type == "field_evaluation":
            st.markdown("### Field Evaluation Rubric")
            st.caption("Official 3-week Field Formative Evaluation Rubric")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("[📄 View Online (PDF)](https://github.com/memari-majid/AI-STER/blob/main/docs/Field_Evaluation_Rubric.pdf)")
            with col2:
                # Add download button
                try:
                    with open("docs/Field_Evaluation_Rubric.pdf", "rb") as pdf_file:
                        st.download_button(
                            label="⬇️ Download PDF",
                            data=pdf_file.read(),
                            file_name="Field_Evaluation_Rubric.pdf",
                            mime="application/pdf"
                        )
                except FileNotFoundError:
                    st.error("PDF file not found")
            with col3:
                st.info("💡 **Tip**: View online or download for offline use")
            
            # Quick reference for Field Evaluation
            st.markdown("""
            **Field Evaluation Components:**
            - 8 Core Competency Items
            - 6 Professional Dispositions (Level 3+ required)
            - Focus on 3-week field experience assessment
            - Competencies must score Level 2+, Dispositions must score Level 3+
            """)
        else:  # STER evaluation
            st.markdown("### STER Evaluation Rubric")
            st.caption("Official USBE STER (Student Teaching Evaluation Rubric) standards")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("[📄 View Online (PDF)](https://github.com/memari-majid/AI-STER/blob/main/docs/STER%20Rubric.pdf)")
            with col2:
                # Add download button
                try:
                    with open("docs/STER Rubric.pdf", "rb") as pdf_file:
                        st.download_button(
                            label="⬇️ Download PDF",
                            data=pdf_file.read(),
                            file_name="STER_Rubric.pdf",
                            mime="application/pdf"
                        )
                except FileNotFoundError:
                    st.error("PDF file not found")
            with col3:
                st.info("💡 **Tip**: View online or download for offline use")
            
            # Quick reference for STER
            st.markdown("""
            **STER Competency Areas:**
            - **Learners and Learning (LL)**: LL1-LL7
            - **Instructional Clarity (IC)**: IC1/IC2, IC3, IC4, IC5/IC6, IC7
            - **Instructional Practice (IP)**: IP1-IP8
            - **Classroom Climate (CC)**: CC1-CC8
            - **Professional Responsibility (PR)**: PR1-PR7
            
            *All 35 competencies require Level 2+ for passing*
            """)
        
        st.divider()
        st.caption("📌 Keep the rubric open in another tab for easy reference during evaluation")
    
    # Get rubric items based on evaluation type
    if rubric_type == "field_evaluation":
        items = get_field_evaluation_items()
        
        # Display evaluation information  
        st.info(f"📋 **Field Evaluation (3-week)** - You will evaluate {len(items)} competencies")
        st.caption("Items based on classroom observation during field experience.")
        
        # Show the breakdown of competencies
        st.markdown("""
        **Competency Areas (8 items total):**
        - **Learners and Learning**: LL3, LL5 (2 items)
        - **Instructional Clarity**: IC1/IC2 (1 item)
        - **Classroom Climate**: CC2, CC3, CC4, CC6, CC8 (5 items)
        """)
    else:  # STER evaluation
        # Get STER items filtered for supervisors (19 competencies)
        all_ster_items = get_ster_items()
        items = filter_items_by_evaluator_role(all_ster_items, 'supervisor')
        
        # Display evaluation information
        st.info(f"📋 **Supervisor STER Evaluation** - You will evaluate {len(items)} competencies")
        st.caption("Items based on classroom observation and lesson planning.")
        
        # Show the breakdown of competencies
        st.markdown("""
        **Competency Areas (19 items total):**
        - **Learners and Learning**: LL2-LL7 (6 items)
        - **Instructional Clarity**: IC1/IC2, IC3, IC4, IC5/IC6, IC7 (5 items)
        - **Instructional Practice**: IP1-IP8 (8 items)
        """)

    
    # Check if we have minimum required information
    if not student_name or not evaluator_name:
        st.warning("Please provide student teacher name and evaluator name to continue.")
        return
    
    # Initialize session state for scores and justifications
    if 'scores' not in st.session_state:
        st.session_state.scores = {}
    if 'justifications' not in st.session_state:
        st.session_state.justifications = {}
    if 'disposition_scores' not in st.session_state:
        st.session_state.disposition_scores = {}
    if 'disposition_comments' not in st.session_state:
        st.session_state.disposition_comments = {}
    if 'observation_notes' not in st.session_state:
        st.session_state.observation_notes = ""
    if 'ai_analyses' not in st.session_state:
        st.session_state.ai_analyses = {}
    if 'current_evaluation_id' not in st.session_state:
        st.session_state.current_evaluation_id = None

    if 'ai_original_data' not in st.session_state:
        st.session_state.ai_original_data = None
    if 'show_ai_comparison' not in st.session_state:
        st.session_state.show_ai_comparison = False
    
    # STEP 3: Classroom Observation Notes
    st.subheader("📝 Step 3: Classroom Observation Notes")
    st.caption("Record your detailed observations of the student teacher's performance during the lesson")
    
    # Input method selection
    col1, col2 = st.columns([3, 1])
    with col1:
        input_method = st.radio(
            "Choose input method:",
            ["✏️ Type/Paste Text", "📤 Upload File"],
            horizontal=True,
            key="observation_input_method"
        )
    
    with col2:
        if input_method == "✏️ Type/Paste Text":
            st.info("💡 Type observations directly")
        else:
            st.info("📄 Upload notes file")
    
    observation_notes = ""
    
    if input_method == "✏️ Type/Paste Text":
        observation_notes = st.text_area(
            "Detailed Observation Notes",
            value=st.session_state.observation_notes,
            height=200,
            placeholder="""Record specific observations about the student teacher's performance:

• How did they introduce the lesson and engage students?
• What teaching strategies and methods were used?
• How did they manage the classroom and respond to student needs?
• What evidence did you see of lesson planning and preparation?
• How did they assess student learning and provide feedback?
• What professional behaviors and dispositions were demonstrated?
• Any specific examples of strengths or areas for improvement?

Be as detailed as possible - these notes will be used to generate evidence-based justifications for each competency area.""",
            help="Detailed observations will help generate more accurate AI justifications",
            key="observation_text_area"
        )
    
    elif input_method == "📤 Upload File":
        uploaded_file = st.file_uploader(
            "Upload Observation Notes File",
            type=['txt', 'pdf', 'docx'],
            help="Upload your observation notes file. Supported formats: Text (.txt), PDF (.pdf), Word (.docx)",
            key="observation_file_uploader"
        )
        
        if uploaded_file is not None:
            # Process the uploaded file
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            try:
                if file_extension == 'txt':
                    # Handle text files
                    file_content = uploaded_file.read().decode('utf-8')
                    st.success(f"✅ File '{uploaded_file.name}' uploaded successfully!")
                    
                    # Preview uploaded content
                    with st.expander("📄 Preview uploaded content", expanded=True):
                        st.text_area(
                            "File Content Preview",
                            value=file_content,
                            height=200,
                            disabled=True,
                            key="file_preview"
                        )
                    
                    # Options to use uploaded content
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        if st.button("📋 Use These Notes", type="primary", key="use_uploaded_notes"):
                            st.session_state.observation_notes = file_content
                            observation_notes = file_content
                            st.success("📝 Uploaded notes have been applied!")
                            st.rerun()
                    
                    with col2:
                        if st.button("➕ Append to Current", key="append_uploaded_notes"):
                            current_notes = st.session_state.observation_notes
                            separator = "\n\n--- Uploaded Notes ---\n\n" if current_notes.strip() else ""
                            combined_notes = current_notes + separator + file_content
                            st.session_state.observation_notes = combined_notes
                            observation_notes = combined_notes
                            st.success("📝 Notes appended successfully!")
                            st.rerun()
                    
                    # Use current session state for display
                    observation_notes = st.session_state.observation_notes
                    
                elif file_extension == 'pdf':
                    # Handle PDF files with automatic text extraction
                    try:
                        import pdfplumber
                        
                        # Extract text using pdfplumber
                        with pdfplumber.open(uploaded_file) as pdf:
                            file_content = ""
                            for page in pdf.pages:
                                text = page.extract_text()
                                if text:
                                    file_content += text + "\n"
                        
                        if file_content.strip():
                            st.success(f"✅ PDF '{uploaded_file.name}' processed successfully!")
                            
                            # Preview uploaded content
                            with st.expander("📄 Preview extracted content", expanded=True):
                                st.text_area(
                                    "Extracted Text Preview",
                                    value=file_content,
                                    height=200,
                                    disabled=True,
                                    key="pdf_preview"
                                )
                            
                            # Options to use uploaded content
                            col1, col2 = st.columns([2, 1])
                            with col1:
                                if st.button("📋 Use These Notes", type="primary", key="use_pdf_notes"):
                                    st.session_state.observation_notes = file_content
                                    observation_notes = file_content
                                    st.success("📝 PDF notes have been applied!")
                                    st.rerun()
                            
                            with col2:
                                if st.button("➕ Append to Current", key="append_pdf_notes"):
                                    current_notes = st.session_state.observation_notes
                                    separator = "\n\n--- Uploaded PDF Notes ---\n\n" if current_notes.strip() else ""
                                    combined_notes = current_notes + separator + file_content
                                    st.session_state.observation_notes = combined_notes
                                    observation_notes = combined_notes
                                    st.success("📝 PDF notes appended successfully!")
                                    st.rerun()
                            
                            # Use current session state for display
                            observation_notes = st.session_state.observation_notes
                        else:
                            st.warning("⚠️ No text could be extracted from this PDF. Please try copying and pasting the content manually.")
                            # Fallback to manual entry
                            observation_notes = st.text_area(
                                "Paste observation notes content:",
                                value=st.session_state.observation_notes,
                                height=200,
                                placeholder="Copy and paste your observation notes from the PDF here...",
                                key="pdf_fallback_text_area"
                            )
                            
                    except ImportError:
                        st.error("❌ PDF processing libraries not available. Please install: pip install pdfplumber")
                        observation_notes = st.session_state.observation_notes
                    except Exception as e:
                        st.error(f"❌ Error processing PDF: {str(e)}")
                        # Fallback to manual entry
                        observation_notes = st.text_area(
                            "Paste observation notes content:",
                            value=st.session_state.observation_notes,
                            height=200,
                            placeholder="Copy and paste your observation notes from the PDF here...",
                            key="pdf_error_fallback_text_area"
                        )
                    
                elif file_extension == 'docx':
                    # Handle DOCX files with automatic text extraction
                    try:
                        from docx import Document
                        
                        # Extract text from Word document with enhanced processing
                        doc = Document(uploaded_file)
                        file_content = ""
                        
                        # Extract text from paragraphs
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                file_content += paragraph.text + "\n"
                        
                        # Extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        file_content += cell.text + " "
                                file_content += "\n"
                        
                        # Extract text from headers and footers
                        for section in doc.sections:
                            if hasattr(section, 'header') and section.header:
                                for paragraph in section.header.paragraphs:
                                    if paragraph.text.strip():
                                        file_content += paragraph.text + "\n"
                            if hasattr(section, 'footer') and section.footer:
                                for paragraph in section.footer.paragraphs:
                                    if paragraph.text.strip():
                                        file_content += paragraph.text + "\n"
                        
                        if file_content.strip():
                            st.success(f"✅ Word document '{uploaded_file.name}' processed successfully!")
                            
                            # Preview uploaded content
                            with st.expander("📄 Preview extracted content", expanded=True):
                                st.text_area(
                                    "Extracted Text Preview",
                                    value=file_content,
                                    height=200,
                                    disabled=True,
                                    key="docx_preview"
                                )
                            
                            # Options to use uploaded content
                            col1, col2 = st.columns([2, 1])
                            with col1:
                                if st.button("📋 Use These Notes", type="primary", key="use_docx_notes"):
                                    st.session_state.observation_notes = file_content
                                    observation_notes = file_content
                                    st.success("📝 Word document notes have been applied!")
                                    st.rerun()
                            
                            with col2:
                                if st.button("➕ Append to Current", key="append_docx_notes"):
                                    current_notes = st.session_state.observation_notes
                                    separator = "\n\n--- Uploaded Word Document Notes ---\n\n" if current_notes.strip() else ""
                                    combined_notes = current_notes + separator + file_content
                                    st.session_state.observation_notes = combined_notes
                                    observation_notes = combined_notes
                                    st.success("📝 Word document notes appended successfully!")
                                    st.rerun()
                            
                            # Use current session state for display
                            observation_notes = st.session_state.observation_notes
                        else:
                            st.warning("⚠️ No text could be extracted from this Word document. Please try copying and pasting the content manually.")
                            # Fallback to manual entry
                            observation_notes = st.text_area(
                                "Paste observation notes content:",
                                value=st.session_state.observation_notes,
                                height=200,
                                placeholder="Copy and paste your observation notes from the Word document here...",
                                key="docx_fallback_text_area"
                            )
                            
                    except ImportError:
                        st.error("❌ Word document processing library not available. Please install: pip install python-docx")
                        observation_notes = st.session_state.observation_notes
                    except Exception as e:
                        st.warning(f"⚠️ Standard processing failed: {str(e)}")
                        
                        # Try alternative extraction method for corrupted files
                        st.info("🔧 Trying alternative extraction method for corrupted/problematic files...")
                        try:
                            import docx2txt
                            uploaded_file.seek(0)  # Reset file pointer
                            file_content = docx2txt.process(uploaded_file)
                            
                            if file_content and file_content.strip():
                                st.success(f"✅ Alternative method successfully extracted text from '{uploaded_file.name}'!")
                                
                                # Preview uploaded content
                                with st.expander("📄 Preview extracted content", expanded=True):
                                    st.text_area(
                                        "Extracted Text Preview",
                                        value=file_content,
                                        height=200,
                                        disabled=True,
                                        key="docx2txt_preview"
                                    )
                                
                                # Options to use uploaded content
                                col1, col2 = st.columns([2, 1])
                                with col1:
                                    if st.button("📋 Use These Notes", type="primary", key="use_docx2txt_notes"):
                                        st.session_state.observation_notes = file_content
                                        observation_notes = file_content
                                        st.success("📝 Extracted notes have been applied!")
                                        st.rerun()
                                
                                with col2:
                                    if st.button("➕ Append to Current", key="append_docx2txt_notes"):
                                        current_notes = st.session_state.observation_notes
                                        separator = "\n\n--- Extracted Word Document Notes ---\n\n" if current_notes.strip() else ""
                                        combined_notes = current_notes + separator + file_content
                                        st.session_state.observation_notes = combined_notes
                                        observation_notes = combined_notes
                                        st.success("📝 Extracted notes appended successfully!")
                                        st.rerun()
                                
                                # Use current session state for display
                                observation_notes = st.session_state.observation_notes
                            else:
                                raise Exception("No text could be extracted")
                                
                        except Exception as e2:
                            st.error(f"❌ Both extraction methods failed: {str(e2)}")
                            st.info("💡 **Troubleshooting tips for corrupted files:**")
                            st.markdown("""
                            - **Font Issues**: Try opening in Word and saving as new .docx file
                            - **Google Docs Method**: Upload to Google Drive → Open with Google Docs → Copy text
                            - **Online Converters**: Try SmallPDF, Zamzar, or CloudConvert
                            - **Document Repair**: In Word, use File → Open → Open and Repair
                            - **Manual Option**: Use the text area below as fallback
                            """)
                            # Fallback to manual entry
                            observation_notes = st.text_area(
                                "Paste observation notes content:",
                                value=st.session_state.observation_notes,
                                height=200,
                                placeholder="Copy and paste your observation notes from the Word document here...",
                                key="docx_error_fallback_text_area"
                            )
                    
                else:
                    st.error(f"❌ Unsupported file type: {file_extension}")
                    observation_notes = st.session_state.observation_notes
                    
            except UnicodeDecodeError:
                st.error("❌ Error reading file. Please ensure it's a valid text file with UTF-8 encoding.")
                observation_notes = st.session_state.observation_notes
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}")
                observation_notes = st.session_state.observation_notes
        else:
            # No file uploaded yet, show current session state
            observation_notes = st.session_state.observation_notes
            if observation_notes:
                st.info("📝 Using previously entered observation notes")
                with st.expander("📄 Current observation notes"):
                    st.text_area(
                        "Current Notes",
                        value=observation_notes,
                        height=150,
                        disabled=True,
                        key="current_notes_preview"
                    )
    
    # Store observation notes in session state
    st.session_state.observation_notes = observation_notes
    
    # STEP 4: Score & Review Competencies (Merged Steps 4, 5, 6)
    st.subheader("🎯 Step 4: Score & Review Competencies")
    if st.session_state.lesson_plan_analysis:
        st.caption("AI analyzes your observation notes and lesson plan, then you score and review each competency")
    else:
        st.caption("AI analyzes your observation notes, then you score and review each competency")
    
    # Generate AI Analysis if not already done
    if observation_notes.strip():
        if openai_service.is_enabled():
            if not st.session_state.get('ai_analyses'):
                col1, col2 = st.columns([2, 1])
                with col1:
                    button_text = "🤖 Generate AI Analysis & Begin Scoring"
                    if not st.session_state.lesson_plan_analysis:
                        button_text += " (Observation Notes Only)"
                        
                    if st.button(button_text, type="primary", key="generate_ai_analysis"):
                        with st.spinner("Analyzing observation notes and generating evidence-based justifications..."):
                            try:
                                # Get lesson plan context if available
                                lesson_plan_context = None
                                if st.session_state.lesson_plan_analysis:
                                    lesson_plan_context = f"Lesson Topic: {st.session_state.lesson_plan_analysis.get('lesson_topic', 'N/A')}\n"
                                    lesson_plan_context += f"Learning Objectives: {', '.join(st.session_state.lesson_plan_analysis.get('learning_objectives', []))}\n"
                                    lesson_plan_context += f"Lesson Structure: {st.session_state.lesson_plan_analysis.get('lesson_structure', 'N/A')}"
                                
                                # Generate AI analysis for all items
                                ai_analyses = openai_service.generate_analysis_for_competencies(
                                    items,
                                    observation_notes,
                                    student_name,
                                    rubric_type,
                                    lesson_plan_context
                                )
                                
                                # Store AI analyses in session state
                                st.session_state.ai_analyses = ai_analyses
                                # Also store as justifications for the save functionality
                                for item_id, analysis in ai_analyses.items():
                                    st.session_state.justifications[item_id] = analysis
                                success_message = f"✅ Generated AI analysis for {len(ai_analyses)} competencies!"
                                if not st.session_state.lesson_plan_analysis:
                                    success_message += " (Based on observation notes only)"
                                st.success(success_message)
                                st.rerun()
                                
                            except Exception as e:
                                st.error(f"Failed to generate AI analysis: {str(e)}")
            
            with col2:
                if st.session_state.get('ai_analyses'):
                    st.success("✅ AI Analysis Complete")
                    st.metric("Competencies Analyzed", len(st.session_state.ai_analyses))
                    
                    if st.button("🔄 Regenerate Analysis", key="regenerate_analysis"):
                        st.session_state.ai_analyses = {}
                        st.session_state.ai_original_data = None
                        st.rerun()
                else:
                    # Show lesson plan status
                    if st.session_state.lesson_plan_analysis:
                        st.info("📄 Lesson plan available")
                    else:
                        st.warning("📄 No lesson plan")
        else:
            st.warning("🤖 AI features disabled. Add OpenAI API key in Settings to enable AI analysis.")
    else:
        st.info("ℹ️ Please add detailed observation notes to enable AI analysis generation.")
        if not st.session_state.lesson_plan_analysis:
            st.warning("⚠️ **Important:** Without a lesson plan, detailed observation notes are crucial for quality AI analysis!")
    
    # Score and Review Section
    if st.session_state.get('ai_analyses') or not openai_service.is_enabled():
        st.markdown("---")
        st.markdown("### 📊 Score and Review Each Competency")
        
        # Check for missing scores
        total_items = len(items)
        scored_items = len([s for s in st.session_state.scores.values() if s is not None])
        missing_scores = total_items - scored_items
        
        if missing_scores > 0:
            st.warning(f"⚠️ **{missing_scores} out of {total_items} competencies still need scores.** Please score all competencies before saving the evaluation.")
        else:
            st.success(f"✅ **All {total_items} competencies have been scored!**")
        
        # Group items by competency area
        competency_groups = {}
        for item in items:
            area = item['competency_area']
            if area not in competency_groups:
                competency_groups[area] = []
            competency_groups[area].append(item)
        
        # Display each competency area
        for area, area_items in competency_groups.items():
            with st.expander(f"📋 **{area}** ({len(area_items)} items)", expanded=True):
                for item in area_items:
                    item_id = item['id']
                    current_score = st.session_state.scores.get(item_id)
                    
                    # Competency header
                    st.markdown(f"### {item['code']}: {item['title']}")
                    st.caption(f"{item['context']}")
                    
                    # Create three columns for integrated interface
                    col1, col2, col3 = st.columns([2, 2, 1])
                    
                    with col1:
                        st.markdown("**📝 Evidence & Justification**")
                        
                        # Get AI analysis or current justification
                        ai_analysis = ""
                        if st.session_state.get('ai_analyses') and item_id in st.session_state.ai_analyses:
                            ai_analysis = st.session_state.ai_analyses[item_id]
                            # Clean warning patterns
                            ai_analysis = ai_analysis.replace('[LIMITED_EVIDENCE]', '').replace('[NO_CONTEXT]', '').replace('[GENERIC]', '').strip()
                        
                        current_justification = st.session_state.justifications.get(item_id, ai_analysis)
                        
                        # Editable text area for justification
                        justification = st.text_area(
                            "Edit justification",
                            value=current_justification,
                            height=120,
                            placeholder="Enter evidence-based justification..." if not ai_analysis else "AI-generated analysis shown - edit as needed",
                            key=f"justification_{item_id}",
                            label_visibility="collapsed",
                            help="Review and edit the AI-generated analysis or write your own justification"
                        )
                        
                        # Update justification in session state
                        if justification.strip():
                            st.session_state.justifications[item_id] = justification
                        elif item_id in st.session_state.justifications and not justification.strip():
                            del st.session_state.justifications[item_id]
                    
                    with col2:
                        st.markdown("**📊 Scoring Rubric**")
                        
                        # Show rubric levels for reference
                        with st.container():
                            for level in ['0', '1', '2', '3']:
                                level_desc = item['levels'].get(level, 'No description')
                                if current_score == int(level):
                                    st.success(f"**Level {level} (Selected):** {level_desc[:100]}...")
                                else:
                                    st.caption(f"**Level {level}:** {level_desc[:100]}...")
                    
                    with col3:
                        st.markdown("**🎯 Score**")
                        
                        def format_score_option(score):
                            if score is None:
                                return "Select..."
                            elif score == "not_observed":
                                return "Not Observed"
                            else:
                                return f"Level {score}"
                        
                        def get_score_index(current_score):
                            if current_score is None:
                                return 0
                            elif current_score == "not_observed":
                                return 1
                            else:
                                return current_score + 2  # Account for None and "not_observed"
                        
                        score = st.selectbox(
                            f"Score for {item['code']}",
                            options=[None, "not_observed", 0, 1, 2, 3],
                            index=get_score_index(current_score),
                            format_func=format_score_option,
                            key=f"score_select_{item_id}",
                            label_visibility="collapsed"
                        )
                        
                        if score is not None:
                            st.session_state.scores[item_id] = score
                            # Auto-populate justification if empty and we have AI analysis
                            if not st.session_state.justifications.get(item_id) and justification.strip():
                                st.session_state.justifications[item_id] = justification
                        
                        # Show score status
                        if score is not None:
                            if score == "not_observed":
                                st.info("👁️ Not Observed")
                            elif isinstance(score, int) and score >= 2:
                                st.success(f"✅ Level {score}")
                            elif isinstance(score, int):
                                st.warning(f"🌱 Level {score}")
                        else:
                            st.error("❌ Not scored")
                        
                        # Individual AI generation button (if needed)
                        if openai_service.is_enabled() and not st.session_state.justifications.get(item_id, "").strip():
                            if st.button("🤖", key=f"gen_{item_id}", help="Generate AI justification for this item"):
                                with st.spinner("Generating..."):
                                    try:
                                        # Get lesson plan context if available
                                        lesson_plan_context = None
                                        if st.session_state.lesson_plan_analysis:
                                            lesson_plan_context = f"Lesson Topic: {st.session_state.lesson_plan_analysis.get('lesson_topic', 'N/A')}"
                                        
                                        ai_justification = openai_service.generate_justification(
                                            item, score if score is not None else 2, student_name, observation_notes, lesson_plan_context
                                        )
                                        st.session_state.justifications[item_id] = ai_justification
                                        if 'ai_analyses' not in st.session_state:
                                            st.session_state.ai_analyses = {}
                                        st.session_state.ai_analyses[item_id] = ai_justification
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"Error: {str(e)}")
                    
                    st.divider()
        
        # Score Summary
        st.markdown("---")
        st.subheader("📈 Scoring Summary")
        
        scored_items = len([s for s in st.session_state.scores.values() if s is not None])
        not_observed_count = len([s for s in st.session_state.scores.values() if s == "not_observed"])
        level_2_plus = len([s for s in st.session_state.scores.values() if isinstance(s, int) and s >= 2])
        level_1_count = len([s for s in st.session_state.scores.values() if s == 1])
        level_0_count = len([s for s in st.session_state.scores.values() if s == 0])
        critical_areas = level_0_count + level_1_count
        
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            if scored_items < len(items):
                st.metric("Items Scored", f"{scored_items}/{len(items)}", delta=f"-{len(items) - scored_items} missing", delta_color="inverse")
            else:
                st.metric("Items Scored", f"{scored_items}/{len(items)}", delta="Complete ✅")
        with col2:
            if scored_items > 0:
                actual_scores = scored_items - not_observed_count  # Exclude not_observed from denominator
                st.metric("Meeting Standards", f"{level_2_plus}/{actual_scores}", help="Level 2+ competencies (excluding not observed)")
            else:
                st.metric("Meeting Standards", "0/0")
        with col3:
            if critical_areas > 0:
                st.metric("Growth Areas", critical_areas, delta="Need Level 2+", delta_color="inverse")
            else:
                st.metric("Growth Areas", 0, delta="All Level 2+ ✅")
        with col4:
            if not_observed_count > 0:
                st.metric("Not Observed", not_observed_count, delta="Review needed", delta_color="off", help="Items marked as not observed during this evaluation")
            else:
                st.metric("Not Observed", 0, delta="All observed ✅")
        with col5:
            all_items_scored = scored_items == len(items)
            meets_minimum = all_items_scored and critical_areas == 0
            if meets_minimum:
                st.metric("Status", "✅ Ready", help="All requirements met")
            elif all_items_scored:
                st.metric("Status", "🌱 Growing", delta=f"{critical_areas} areas", delta_color="inverse")
            else:
                st.metric("Status", "⏳ Incomplete")
        
        # Not Observed Items Warning
        if not_observed_count > 0:
            st.warning(f"⚠️ **{not_observed_count} competencies marked as 'Not Observed'** - These items were not demonstrated during this observation session. Consider scheduling additional observations to evaluate these competencies if they are required for completion.")
            
            # Show which items are not observed
            not_observed_items = [item for item in items if st.session_state.scores.get(item['id']) == "not_observed"]
            if not_observed_items:
                st.markdown("**Items marked as 'Not Observed':**")
                for item in not_observed_items:
                    st.write(f"• **{item['code']}**: {item['title']}")
                st.info("💡 **Tip**: Consider if these competencies can be evaluated in future observations or through other assessment methods.")
        
        # Growth Areas Alert
        if scored_items > 0 and critical_areas > 0:
            st.markdown("---")
            st.subheader("🌟 Areas for Growth")
            
            level_0_items = [item for item in items if st.session_state.scores.get(item['id']) == 0]
            level_1_items = [item for item in items if st.session_state.scores.get(item['id']) == 1]
            
            if level_0_items:
                st.error("**Level 0 - Priority Development Areas:**")
                for item in level_0_items:
                    st.write(f"• {item['code']}: {item['title']}")
            
            if level_1_items:
                st.warning("**Level 1 - Approaching Competency:**")
                for item in level_1_items:
                    st.write(f"• {item['code']}: {item['title']}")

    # STEP 5: Professional Dispositions (Field Evaluation Only)  
    if rubric_type == "field_evaluation":
        st.subheader("🌟 Step 5: Professional Dispositions (Supervisor Manual Assessment)")
        st.caption("Supervisors evaluate dispositions based on observations across multiple interactions and timeframes, not limited to a single lesson.")
        
        # Add context about manual assessment
        st.info("""
        **Professional Dispositions Assessment:**
        • ⏰ **Broader Time Scope**: Based on multiple observations, conferences, and interactions over time
        • 👨‍🏫 **Supervisor Assessment**: Manual evaluation by supervisor (no AI assistance)  
        • 📝 **Feedback Required**: Provide specific comments and suggestions for growth
        • ⭐ **Completion Requirement**: All dispositions must score Level 3+ to complete evaluation
        """)
        
        st.markdown("**Score Each Disposition (Level 1-4) and Provide Feedback**")
        
        all_dispositions_scored = True
        
        for disposition in dispositions:
            disp_id = disposition['id']
            current_score = st.session_state.disposition_scores.get(disp_id)
            current_comment = st.session_state.disposition_comments.get(disp_id, "")
            
            # Create header for each disposition
            st.markdown(f"**{disposition['name']}**")
            st.caption(disposition['description'])
            
            # Create scoring and comment interface for each disposition  
            col1, col2 = st.columns([1, 2])
            
            with col1:
                score = st.selectbox(
                    f"Score",
                    options=[None, 1, 2, 3, 4],
                    index=0 if current_score is None else current_score,
                    format_func=lambda x: "Select..." if x is None else f"Level {x} - {get_disposition_level_name(x)}",
                    key=f"disposition_select_{disp_id}",
                    help=f"Select score level for {disposition['name']}"
                )
                
                if score is not None:
                    st.session_state.disposition_scores[disp_id] = score
                else:
                    all_dispositions_scored = False
            
            with col2:
                comment = st.text_area(
                    "Feedback & Suggestions",
                    value=current_comment,
                    height=100,
                    placeholder="Provide specific feedback, observations, and suggestions for professional growth...",
                    key=f"disposition_comment_{disp_id}",
                    help="Enter detailed feedback based on your observations across multiple interactions",
                    max_chars=1000
                )
                
                # Update session state for comments
                st.session_state.disposition_comments[disp_id] = comment
                
                # Character count
                st.caption(f"{len(comment)}/1000 characters")
            
            # Show selected score description and color coding
            display_score = score if score is not None else current_score
            if display_score is not None:
                score_desc = get_disposition_level_name(display_score)
                if display_score >= 3:
                    st.success(f"**Level {display_score}:** {score_desc}")
                elif display_score == 2:
                    st.warning(f"**Level {display_score}:** {score_desc} (Level 3+ required)")
                else:
                    st.warning(f"**Level {display_score}:** {score_desc} (Level 3+ required)")
            else:
                # Show warning for unscored disposition
                st.warning(f"⏸️ **Not yet scored** - {disposition['name']} requires a score level to complete the evaluation")
            
            st.divider()
    
    # AI Analysis
    if openai_service.is_enabled() and st.session_state.scores:
        st.subheader("🤖 Overall Performance Analysis")
        st.caption("AI analysis focused on specific areas needing improvement and actionable guidance for growth")
        
        # Show preview of areas needing attention
        level_1_items = [item_id for item_id, score in st.session_state.scores.items() if score == 1]
        low_dispositions = []
        if rubric_type == "field_evaluation":
            low_dispositions = [disp_id for disp_id, score in st.session_state.disposition_scores.items() if score < 3]
        
        if level_1_items or low_dispositions:
            warning_msg = f"⚠️ **Areas Requiring Attention**: {len(level_1_items)} competencies at Level 1"
            if low_dispositions:
                warning_msg += f", {len(low_dispositions)} dispositions below Level 3"
            st.warning(warning_msg)
        else:
            st.success("✅ **All areas meeting minimum requirements** - Analysis will focus on strengths and growth opportunities")
        
        if st.button("Generate Targeted Improvement Analysis"):
            with st.spinner("Analyzing specific areas needing improvement..."):
                try:
                    # Enhanced analysis that focuses on specific improvement areas
                    analysis = openai_service.analyze_evaluation(
                        st.session_state.scores,
                        st.session_state.justifications,
                        st.session_state.disposition_scores,
                        rubric_type
                    )
                    
                    # Store analysis in session state
                    st.session_state.targeted_improvement_analysis = analysis
                    
                    st.success("**Targeted Performance Analysis:**")
                    st.info(analysis)
                    
                    # Updated guidance reflecting new approach
                    st.markdown("""
                    **How to use this analysis:**
                    • 🎯 **Focus on Level 1 areas**: Students must achieve Level 2+ in ALL competencies to pass
                    • 📝 **Use specific improvement steps**: Provide targeted guidance for moving from Level 1 to Level 2
                    • 🗣️ **Conference planning**: Address specific concerns rather than overall averages
                    • 📈 **Track progress**: Monitor improvement in identified growth areas
                    • 💪 **Leverage strengths**: Use Level 2-3 areas to support growth in struggling competencies
                    """)
                    
                except Exception as e:
                    st.error(f"AI analysis failed: {str(e)}")
        
        # Display existing analysis if available
        elif st.session_state.get('targeted_improvement_analysis'):
            st.success("**Targeted Performance Analysis:**")
            st.info(st.session_state.targeted_improvement_analysis)
            
            # Allow regeneration
            if st.button("🔄 Regenerate Analysis"):
                st.session_state.targeted_improvement_analysis = None
                st.rerun()
    
    # Save buttons
    st.subheader("Save Evaluation")
    
    # Check for completion before saving
    missing_competency_scores = len(items) - len([s for s in st.session_state.scores.values() if s is not None])
    missing_disposition_scores = 0
    if rubric_type == "field_evaluation":
        missing_disposition_scores = len(dispositions) - len([s for s in st.session_state.disposition_scores.values() if s is not None])
    
    if missing_competency_scores > 0 or missing_disposition_scores > 0:
        warning_msg = "⚠️ **Incomplete Evaluation Warning:**\n"
        if missing_competency_scores > 0:
            warning_msg += f"\n• {missing_competency_scores} competencies still need scores"
        if missing_disposition_scores > 0:
            warning_msg += f"\n• {missing_disposition_scores} professional dispositions still need scores"
        warning_msg += "\n\n*You can save as draft with missing scores, but completion requires all items to be scored.*"
        st.warning(warning_msg)
    else:
        if rubric_type == "field_evaluation":
            st.success("✅ **Evaluation is complete!** All competencies and dispositions have been scored.")
        else:
            st.success("✅ **Evaluation is complete!** All competencies have been scored.")
    
    # Calculate total score
    total_score = sum(score for score in st.session_state.scores.values() if isinstance(score, int))
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("💾 Save as Draft"):
            # Get extracted info for additional fields
            extracted_info = st.session_state.get('extracted_info', {})
            
            # Generate or use existing evaluation ID
            eval_id = st.session_state.get('current_evaluation_id', str(uuid.uuid4()))
            st.session_state.current_evaluation_id = eval_id
            
            evaluation = {
                'id': eval_id,
                'student_name': student_name,
                'evaluator_name': evaluator_name,
                'evaluator_role': evaluator_role,
                'rubric_type': rubric_type,
                'evaluated_items_count': len(items),  # Track how many items this role evaluated
                'scores': st.session_state.scores,
                'justifications': st.session_state.justifications,
                'disposition_scores': st.session_state.disposition_scores if rubric_type == "field_evaluation" else {},
                'disposition_comments': st.session_state.disposition_comments if rubric_type == "field_evaluation" else {},
                'total_score': total_score,
                'status': 'draft',
                'created_at': datetime.now().isoformat(),
                'lesson_plan_provided': st.session_state.lesson_plan_analysis is not None,
                'lesson_plan_method': input_method if 'input_method' in locals() else 'unknown',
                'ai_analyses': st.session_state.get('ai_analyses', {}),
                'targeted_improvement_analysis': st.session_state.get('targeted_improvement_analysis', ''),
                # Dashboard fields
                'subject_area': extracted_info.get('subject_area', ''),
                'department': extracted_info.get('department', 'Secondary'),
                'semester': extracted_info.get('semester', 'Spring 2025'),
                'grade_levels': extracted_info.get('grade_levels', ''),
                'school_name': extracted_info.get('school_name', ''),
                'class_size': extracted_info.get('class_size', 20)
            }
            
            # Add AI original data if available
            if st.session_state.get('ai_original_data'):
                evaluation['ai_original'] = st.session_state.ai_original_data
                evaluation['has_ai_original'] = True
                evaluation['ai_original_saved_at'] = st.session_state.ai_original_data.get('saved_at')
            
            save_evaluation(evaluation)
            st.success("Evaluation saved as draft!")
    
    with col2:
        if st.button("✅ Complete Evaluation"):
            # Validation
            if rubric_type == "field_evaluation":
                # For Field Evaluation, validate with dispositions
                errors = validate_evaluation(
                    st.session_state.scores,
                    st.session_state.justifications,
                    st.session_state.disposition_scores,
                    items,
                    dispositions
                )
            else:
                # For STER, validate without dispositions
                errors = validate_evaluation(
                    st.session_state.scores,
                    st.session_state.justifications,
                    {},  # Empty dispositions for STER
                    items,
                    []   # Empty dispositions list for STER
                )
            
            if errors:
                st.warning("⚠️ **Evaluation has issues:**")
                for error in errors:
                    st.error(error)
                st.info("The evaluation will be saved with a 'needs_improvement' status.")
            
            # Save evaluation regardless of validation errors
            # Get extracted info for additional fields
            extracted_info = st.session_state.get('extracted_info', {})
            
            # Generate or use existing evaluation ID
            eval_id = st.session_state.get('current_evaluation_id', str(uuid.uuid4()))
            st.session_state.current_evaluation_id = eval_id
            
            evaluation = {
                'id': eval_id,
                'student_name': student_name,
                'evaluator_name': evaluator_name,
                'evaluator_role': evaluator_role,
                'rubric_type': rubric_type,
                'scores': st.session_state.scores,
                'justifications': st.session_state.justifications,
                'disposition_scores': st.session_state.disposition_scores,
                'disposition_comments': st.session_state.disposition_comments,
                'total_score': total_score,
                'status': 'needs_improvement' if errors else 'completed',
                'validation_errors': errors if errors else [],
                'created_at': datetime.now().isoformat(),
                'completed_at': datetime.now().isoformat(),
                'lesson_plan_provided': st.session_state.lesson_plan_analysis is not None,
                'lesson_plan_method': input_method if 'input_method' in locals() else 'unknown',
                'ai_analyses': st.session_state.get('ai_analyses', {}),
                'targeted_improvement_analysis': st.session_state.get('targeted_improvement_analysis', ''),
                # Dashboard fields
                'subject_area': extracted_info.get('subject_area', ''),
                'department': extracted_info.get('department', 'Secondary'),
                'semester': extracted_info.get('semester', 'Spring 2025'),
                'grade_levels': extracted_info.get('grade_levels', ''),
                'school_name': extracted_info.get('school_name', ''),
                'class_size': extracted_info.get('class_size', 20)
            }
            
            # Add AI original data if available
            if st.session_state.get('ai_original_data'):
                evaluation['ai_original'] = st.session_state.ai_original_data
                evaluation['has_ai_original'] = True
                evaluation['ai_original_saved_at'] = st.session_state.ai_original_data.get('saved_at')
            
            save_evaluation(evaluation)
            
            if errors:
                st.warning("🎯 **Evaluation saved with 'needs improvement' status**")
                st.caption("The student teacher requires additional support in the identified areas.")
            else:
                st.success("🎉 Evaluation completed successfully!")
            
            # Quick PDF Download for Completed Evaluation
            st.markdown("---")
            st.subheader("📄 Download Completed Evaluation Report")
            col1_quick, col2_quick = st.columns(2)
            
            with col1_quick:
                # Prepare data for PDF generation
                pdf_data = {
                    'rubric_type': rubric_type,
                    'student_name': student_name,
                    'evaluator_name': evaluator_name,
                    'date': datetime.now().strftime('%Y-%m-%d'),
                    'school': st.session_state.get('school_name', 'N/A'),
                    'subject': st.session_state.get('subject_grade', 'N/A'),
                    'is_formative': True,
                    'competency_scores': [],
                    'total_items': len(items),
                    'meeting_expectations': sum(1 for score in st.session_state.scores.values() if isinstance(score, int) and score >= 2),
                    'areas_for_growth': sum(1 for score in st.session_state.scores.values() if isinstance(score, int) and score < 2),
                    'targeted_improvement_analysis': st.session_state.get('targeted_improvement_analysis', ''),
                    'status': 'needs_improvement' if errors else 'completed'
                }
                
                # Add competency scores
                for item in items:
                    item_id = item['id']
                    score = st.session_state.scores.get(item_id, None)
                    pdf_data['competency_scores'].append({
                        'competency': f"{item['code']}: {item['title']}",
                        'score': score if score is not None else 'Not Scored',
                        'justification': st.session_state.justifications.get(item_id, 'No justification provided')
                    })
                
                # Add dispositions for field evaluations
                if rubric_type == 'field_evaluation' and st.session_state.disposition_scores:
                    pdf_data['dispositions'] = []
                    for disposition in dispositions:
                        disp_id = disposition['id']
                        score = st.session_state.disposition_scores.get(disp_id, None)
                        pdf_data['dispositions'].append({
                            'disposition': disposition['name'],
                            'score': score if score is not None else 'Not Scored',
                            'notes': st.session_state.disposition_comments.get(disp_id, '')
                        })
                
                # Add AI analysis if available
                ai_analyses = st.session_state.get('ai_analyses', {})
                if ai_analyses:
                    all_analyses = []
                    strengths = []
                    areas_for_growth = []
                    
                    for item in items:
                        analysis = ai_analyses.get(item['id'])
                        if analysis:
                            item_name = f"{item['code']}: {item['title']}"
                            all_analyses.append(f"{item_name}\n{analysis}")
                            
                            if any(word in analysis.lower() for word in ['strong', 'excellent', 'effective', 'well', 'good', 'demonstrates']):
                                strengths.append(f"{item['code']}: {analysis[:150]}...")
                            
                            if any(word in analysis.lower() for word in ['improve', 'develop', 'consider', 'could', 'should', 'needs']):
                                areas_for_growth.append(f"{item['code']}: {analysis[:150]}...")
                    
                    pdf_data['ai_analysis'] = {
                        'strengths': strengths[:5],
                        'areas_for_growth': areas_for_growth[:5],
                        'recommendations': [],
                        'full_analyses': all_analyses
                    }
                
                # Generate PDF
                try:
                    pdf_bytes = pdf_service.generate_evaluation_pdf(pdf_data)
                    
                    # Create filename
                    status_text = "NEEDS_IMPROVEMENT" if errors else "COMPLETED"
                    filename = f"{student_name.replace(' ', '_')}_{rubric_type}_{status_text}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                    
                    st.download_button(
                        label="📄 Download Final Report (PDF)",
                        data=pdf_bytes,
                        file_name=filename,
                        mime="application/pdf",
                        help="Download the completed evaluation report",
                        type="primary"
                    )
                    
                except Exception as e:
                    st.error(f"Error generating PDF: {str(e)}")
            
            with col2_quick:
                if errors:
                    st.info("📋 **Report Status:** Needs Improvement\nThis report includes identified areas requiring additional support.")
                else:
                    st.success("📋 **Report Status:** Complete\nThis report shows a fully completed evaluation.")
                
                st.info("💡 **Tip:** You can also use the 'Download Draft Report' section below for additional options.")
            
            
    # PDF Download Section - Always Available
    st.markdown("---")
    st.subheader("📄 Download Draft Report")
    st.caption("Generate a PDF report at any stage of the evaluation process")
    
    col1_pdf, col2_pdf = st.columns(2)
    
    with col1_pdf:
        # Check if there's any data to export
        has_scores = len(st.session_state.scores) > 0
        has_justifications = len(st.session_state.justifications) > 0
        has_basic_info = student_name.strip() != "" and evaluator_name.strip() != ""
        
        if has_basic_info:
            # Prepare data for PDF generation
            pdf_data = {
                'rubric_type': rubric_type,
                'student_name': student_name,
                'evaluator_name': evaluator_name,
                'date': datetime.now().strftime('%Y-%m-%d'),
                'school': st.session_state.get('school_name', 'N/A'),
                'subject': st.session_state.get('subject_grade', 'N/A'),
                'is_formative': True,
                'competency_scores': [],
                'total_items': len(items),
                'meeting_expectations': sum(1 for score in st.session_state.scores.values() if isinstance(score, int) and score >= 2),
                'areas_for_growth': sum(1 for score in st.session_state.scores.values() if isinstance(score, int) and score < 2),
                'targeted_improvement_analysis': st.session_state.get('targeted_improvement_analysis', ''),
                'status': 'draft'  # Mark as draft since it's always available
            }
            
            # Add competency scores (including None values for incomplete items)
            for item in items:
                item_id = item['id']
                score = st.session_state.scores.get(item_id, None)
                pdf_data['competency_scores'].append({
                    'competency': f"{item['code']}: {item['title']}",
                    'score': score if score is not None else 'Not Scored',
                    'justification': st.session_state.justifications.get(item_id, 'No justification provided')
                })
            
            # Add dispositions for field evaluations
            if rubric_type == 'field_evaluation' and st.session_state.disposition_scores:
                pdf_data['dispositions'] = []
                for disposition in dispositions:
                    disp_id = disposition['id']
                    score = st.session_state.disposition_scores.get(disp_id, None)
                    pdf_data['dispositions'].append({
                        'disposition': disposition['name'],
                        'score': score if score is not None else 'Not Scored',
                        'notes': st.session_state.disposition_comments.get(disp_id, '')
                    })
            
            # Add AI analysis if available
            ai_analyses = st.session_state.get('ai_analyses', {})
            if ai_analyses:
                # Collect all AI analyses
                all_analyses = []
                strengths = []
                areas_for_growth = []
                
                for item in items:
                    analysis = ai_analyses.get(item['id'])
                    if analysis:
                        item_name = f"{item['code']}: {item['title']}"
                        all_analyses.append(f"{item_name}\n{analysis}")
                        
                        # Extract strengths and growth areas
                        if any(word in analysis.lower() for word in ['strong', 'excellent', 'effective', 'well', 'good', 'demonstrates']):
                            strengths.append(f"{item['code']}: {analysis[:150]}...")
                        
                        if any(word in analysis.lower() for word in ['improve', 'develop', 'consider', 'could', 'should', 'needs']):
                            areas_for_growth.append(f"{item['code']}: {analysis[:150]}...")
                
                pdf_data['ai_analysis'] = {
                    'strengths': strengths[:5],
                    'areas_for_growth': areas_for_growth[:5],
                    'recommendations': [],
                    'full_analyses': all_analyses
                }
            
            # Generate PDF
            try:
                pdf_bytes = pdf_service.generate_evaluation_pdf(pdf_data)
                
                # Create filename with draft indicator
                status_indicator = "DRAFT" if (len(st.session_state.scores) < len(items)) else "COMPLETE"
                filename = f"{student_name.replace(' ', '_')}_{rubric_type}_{status_indicator}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                
                col_dl1, col_dl2, col_dl3 = st.columns(3)
                
                with col_dl1:
                    st.download_button(
                        label="📄 Download Current Report (PDF)",
                        data=pdf_bytes,
                        file_name=filename,
                        mime="application/pdf",
                        help="Download the evaluation report in its current state"
                    )
                
                with col_dl2:
                    # Store AI original data automatically when AI analysis is available
                    if st.session_state.get('ai_analyses') and not st.session_state.get('ai_original_data'):
                        st.session_state.ai_original_data = {
                            'justifications': st.session_state.get('justifications', {}).copy(),
                            'ai_analyses': st.session_state.get('ai_analyses', {}).copy(),
                            'scores': st.session_state.get('scores', {}).copy(),
                            'observation_notes': st.session_state.get('observation_notes', ''),
                            'lesson_plan_analysis': st.session_state.get('lesson_plan_analysis', None),
                            'saved_at': datetime.now().isoformat()
                        }
                    
                    # Empty column for spacing (removed Save AI Version button)
                    st.empty()
                
                with col_dl3:
                    # AI Performance Evaluation button - shows comparison
                    if st.session_state.get('ai_analyses'):
                        if st.button("🤖 AI Performance Evaluation", key="ai_performance_eval",
                                   help="View comparison between AI-generated and supervisor-revised feedback"):
                            st.session_state.show_ai_comparison = True
                            st.rerun()
                    else:
                        st.button("🤖 AI Performance Evaluation", key="ai_performance_eval_disabled",
                                 disabled=True,
                                 help="Generate AI analysis first to enable comparison")
                
                # Show status information
                if len(st.session_state.scores) < len(items):
                    missing_count = len(items) - len(st.session_state.scores)
                    st.info(f"ℹ️ This is a **DRAFT** report. {missing_count} competencies still need scores.")
                else:
                    st.success("✅ This report includes all competency scores.")
                    
            except Exception as e:
                st.error(f"Error generating PDF: {str(e)}")
        else:
            st.warning("⚠️ Please fill in student name and evaluator name to enable PDF download.")
    
    with col2_pdf:
        if has_basic_info:
            st.info("💡 **Tip**: You can download a report at any stage. Draft reports will be clearly marked.")
            
            # Progress indicator
            total_items = len(items)
            scored_items = len([s for s in st.session_state.scores.values() if s is not None])
            progress = scored_items / total_items if total_items > 0 else 0
            
            st.metric("Progress", f"{scored_items}/{total_items}", f"{progress:.0%}")
            
            if rubric_type == "field_evaluation":
                total_dispositions = len(dispositions) 
                scored_dispositions = len([s for s in st.session_state.disposition_scores.values() if s is not None])
                st.metric("Dispositions", f"{scored_dispositions}/{total_dispositions}")
        else:
            st.info("Fill in the basic information above to enable PDF download.")
    
    # AI Performance Evaluation Comparison View
    if st.session_state.get('show_ai_comparison', False):
        st.markdown("---")
        st.subheader("🤖 AI Performance Evaluation - Comparison Report")
        st.caption("Showing AI-generated content vs. Supervisor revisions")
        
        # Close button
        if st.button("✖️ Close Comparison", key="close_comparison_form"):
            st.session_state.show_ai_comparison = False
            st.rerun()
        
        if st.session_state.get('ai_original_data'):
            ai_data = st.session_state.ai_original_data
            current_data = {
                'justifications': st.session_state.get('justifications', {}),
                'scores': st.session_state.get('scores', {}),
                'observation_notes': st.session_state.get('observation_notes', '')
            }
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                ai_saved_time = ai_data.get('saved_at', 'Unknown')
                st.metric("AI Version Saved", ai_saved_time[:19] if len(ai_saved_time) > 19 else ai_saved_time)
            with col2:
                total_items = len(items)
                st.metric("Total Competencies", total_items)
            with col3:
                modified_count = sum(1 for item_id in ai_data.get('justifications', {})
                                   if ai_data['justifications'].get(item_id, '') != current_data['justifications'].get(item_id, ''))
                st.metric("Modified Justifications", modified_count)
            with col4:
                score_changes = sum(1 for item_id in ai_data.get('scores', {})
                                  if ai_data['scores'].get(item_id) != current_data['scores'].get(item_id))
                st.metric("Score Changes", score_changes)
            
            # Detailed comparison for each competency
            st.markdown("### 📋 Competency-by-Competency Comparison")
            
            for item in items:
                item_id = item['id']
                if item_id in ai_data.get('justifications', {}):
                    with st.expander(f"{item['code']}: {item['title']}", expanded=False):
                        col1_comp, col2_comp = st.columns(2)
                        
                        with col1_comp:
                            st.markdown("**🤖 AI Original:**")
                            ai_score = ai_data.get('scores', {}).get(item_id, 'Not scored')
                            st.info(f"Score: {ai_score}")
                            st.write(ai_data['justifications'].get(item_id, 'No justification'))
                        
                        with col2_comp:
                            st.markdown("**✏️ Supervisor Revised:**")
                            current_score = current_data['scores'].get(item_id, 'Not scored')
                            if current_score != ai_score:
                                st.warning(f"Score: {current_score} (changed from {ai_score})")
                            else:
                                st.success(f"Score: {current_score} (unchanged)")
                            
                            current_just = current_data['justifications'].get(item_id, 'No justification')
                            if current_just != ai_data['justifications'].get(item_id):
                                st.markdown("📝 _Modified justification:_")
                            st.write(current_just)
            
            # Export comparison data
            st.markdown("### 📊 Export Comparison Data")
            comparison_data = {
                'evaluation_info': {
                    'student_name': student_name,
                    'evaluator_name': evaluator_name,
                    'evaluation_date': evaluation_date.isoformat() if 'evaluation_date' in locals() else datetime.now().isoformat(),
                    'rubric_type': rubric_type
                },
                'ai_original': ai_data,
                'supervisor_final': current_data,
                'comparison_summary': {
                    'total_competencies': len(items),
                    'modified_justifications': modified_count,
                    'score_changes': score_changes,
                    'ai_version_saved_at': ai_data.get('saved_at', 'Unknown')
                },
                'export_timestamp': datetime.now().isoformat()
            }
            
            # Generate PDF comparison report
            try:
                comparison_data['items'] = items  # Add items for PDF generation
                comparison_pdf_bytes = pdf_service.generate_comparison_pdf(comparison_data)
                
                st.download_button(
                    "📥 Download Comparison Report (PDF)",
                    comparison_pdf_bytes,
                    f"AI_Comparison_{student_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    "application/pdf",
                    help="Download the comparison report as a formatted PDF"
                )
            except Exception as e:
                st.error(f"Error generating comparison PDF: {str(e)}")
                # Fallback to JSON
                comparison_json = json.dumps(comparison_data, indent=2)
                st.download_button(
                    "📥 Download Comparison Data (JSON)",
                    comparison_json,
                    f"ai_comparison_{student_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    "application/json",
                    help="Download the comparison data as JSON (fallback)"
                )

def show_test_data():
    """Simple demo and visualization page"""
    st.header("🧪 Demo & Visualizations")
    
    # Simple two-column layout for quick actions
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 🚀 Quick Start")
        st.info("Click the button below to generate demo data and visualizations")
        
        if st.button("📊 Generate Demo Data & Charts", type="primary", use_container_width=True):
            with st.spinner("Creating demo scenarios and visualizations..."):
                # Generate everything automatically
                try:
                    # Import at the top of the function to check early
                    import matplotlib.pyplot as plt
                    import numpy as np
                    
                    # Load or generate demo data
                    evaluations = load_evaluations()
                    if len(evaluations) < 5:
                        st.info("Loading demo scenarios...")
                        os.system("python load_demo_data.py")
                        evaluations = load_evaluations()
                    
                    # Generate all visualizations
                    st.success("✅ Demo data loaded! Generating visualizations...")
                    
                    # Store in session state
                    st.session_state.demo_ready = True
                    st.session_state.evaluations = evaluations
                    st.rerun()
                    
                except ImportError as e:
                    st.error(f"Missing required package: {e}")
                    st.info("Run: pip install matplotlib seaborn")
    
    with col2:
        st.markdown("### 📋 What You'll Get")
        st.markdown("""
        - 5 realistic demo scenarios
        - Performance distribution chart
        - AI impact comparison
        - Time savings visualization
        - Subject area breakdown
        """)
    
    # If demo is ready, show everything
    if st.session_state.get('demo_ready', False) and st.session_state.get('evaluations'):
        st.markdown("---")
        st.success("✅ Demo Ready! Here are your visualizations:")
        
        # Create simple tabs
        tab1, tab2, tab3 = st.tabs(["📊 Charts for Poster", "👥 Demo Scenarios", "📸 Export Graphics"])
        
        with tab1:
            st.markdown("### 📊 Research-Based Visualizations")
            
            # Option to choose visualization type
            viz_type = st.radio(
                "Select visualization category:",
                ["Quick Summary", "Detailed Research Metrics", "Generate All"],
                horizontal=True
            )
            
            if viz_type == "Quick Summary":
                # Show the most important charts for poster
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### Time & Efficiency Impact")
                    # Import and generate
                    import matplotlib.pyplot as plt
                    import numpy as np
                    
                    fig, ax = plt.subplots(figsize=(8, 6))
                    
                    # Time comparison based on research
                    activities = ['Document\nReview', 'Evidence\nExtraction', 'Justification\nWriting', 'Score\nCalc', 'Report\nGen']
                    manual_time = [8, 12, 18, 3, 4]  # Total: 45 minutes
                    ai_time = [2, 1, 8, 1, 3]       # Total: 15 minutes
                    
                    x = np.arange(len(activities))
                    width = 0.35
                    
                    bars1 = ax.bar(x - width/2, manual_time, width, label='Manual (45 min)', color='#AAAAAB')
                    bars2 = ax.bar(x + width/2, ai_time, width, label='AI-Assisted (15 min)', color='#185C33')
                    
                    ax.set_xlabel('Task')
                    ax.set_ylabel('Time (minutes)')
                    ax.set_title('Time Allocation by Task', fontsize=14, fontweight='bold')
                    ax.set_xticks(x)
                    ax.set_xticklabels(activities, rotation=45, ha='right')
                    ax.legend()
                    
                    # Add time labels
                    for bars in [bars1, bars2]:
                        for bar in bars:
                            height = bar.get_height()
                            ax.text(bar.get_x() + bar.get_width()/2., height,
                                   f'{int(height)}', ha='center', va='bottom')
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    plt.close()
                
                with col2:
                    st.markdown("#### Evidence Quality Improvement")
                    
                    fig, ax = plt.subplots(figsize=(8, 6))
                    
                    # Evidence extraction improvements
                    evidence_types = ['Specific\nExamples', 'Standards\nAlignment', 'Student\nImpact', 'Teaching\nStrategies']
                    manual_evidence = [2.1, 1.8, 1.5, 2.3]
                    ai_evidence = [5.4, 4.8, 4.2, 5.1]
                    
                    x = np.arange(len(evidence_types))
                    
                    bars1 = ax.bar(x - width/2, manual_evidence, width, label='Manual', color='#AAAAAB')
                    bars2 = ax.bar(x + width/2, ai_evidence, width, label='AI-Assisted', color='#185C33')
                    
                    ax.set_ylabel('Average Evidence Items')
                    ax.set_title('Evidence Extraction Quality', fontsize=14, fontweight='bold')
                    ax.set_xticks(x)
                    ax.set_xticklabels(evidence_types)
                    ax.legend()
                    
                    # Add improvement factors
                    for i, (manual, ai) in enumerate(zip(manual_evidence, ai_evidence)):
                        factor = ai / manual
                        ax.text(x[i], ai + 0.2, f'{factor:.1f}x', ha='center', 
                               fontweight='bold', color='#78BE3F')
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    plt.close()
                
                # Key Research Findings
                st.markdown("### 🔬 Key Research Findings")
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Time Reduction", "67%", "30 min saved per evaluation")
                with col2:
                    st.metric("Evidence Increase", "172%", "3.2 → 8.7 items extracted")
                with col3:
                    st.metric("Consistency Score", "31%", "72 → 94 points improvement")
                with col4:
                    st.metric("Feedback Length", "3x", "~45 → ~130 words")
                
                # Semester impact
                st.markdown("### 📅 Semester Impact Analysis")
                
                fig, ax = plt.subplots(figsize=(12, 6))
                
                weeks = np.arange(1, 17)
                evals_per_week = 5
                manual_hours = weeks * evals_per_week * 45 / 60
                ai_hours = weeks * evals_per_week * 15 / 60
                
                ax.plot(weeks, manual_hours, 'o-', color='#AAAAAB', linewidth=2, markersize=6, label='Manual Process')
                ax.plot(weeks, ai_hours, 's-', color='#185C33', linewidth=2, markersize=6, label='AI-Assisted')
                ax.fill_between(weeks, ai_hours, manual_hours, alpha=0.3, color='#78BE3F', label='Time Saved')
                
                ax.set_xlabel('Weeks in Semester')
                ax.set_ylabel('Cumulative Hours')
                ax.set_title('Cumulative Time Investment Over One Semester', fontsize=14, fontweight='bold')
                ax.legend()
                ax.grid(True, alpha=0.3)
                
                # Add annotation
                total_saved = manual_hours[-1] - ai_hours[-1]
                ax.annotate(f'{total_saved:.0f} hours saved\nper semester',
                           xy=(12, manual_hours[11]), xytext=(10, manual_hours[11] + 10),
                           arrowprops=dict(arrowstyle='->', color='#B45336'),
                           bbox=dict(boxstyle="round,pad=0.5", facecolor='#78BE3F', alpha=0.8),
                           fontsize=12, ha='center')
                
                plt.tight_layout()
                st.pyplot(fig)
                plt.close()
            
            elif viz_type == "Detailed Research Metrics":
                st.info("📈 Generating comprehensive research visualizations...")
                
                # Run the research visualization generator
                with st.spinner("Creating research-based charts..."):
                    from generate_research_visualizations import create_all_research_visualizations
                    create_all_research_visualizations()
                
                # Display the generated images
                st.success("✅ Research visualizations generated!")
                
                # Show generated images in a grid
                col1, col2 = st.columns(2)
                
                research_images = [
                    ("research_quality_comparison.png", "Evaluation Quality Improvements"),
                    ("research_performance_analysis.png", "Performance Level Analysis"),
                    ("research_workflow_efficiency.png", "Workflow Efficiency Analysis"),
                    ("research_feedback_quality.png", "Feedback Quality Analysis")
                ]
                
                for i, (filename, title) in enumerate(research_images):
                    with col1 if i % 2 == 0 else col2:
                        if os.path.exists(filename):
                            st.markdown(f"#### {title}")
                            st.image(filename, use_column_width=True)
                
                # Show comprehensive summary separately
                if os.path.exists("research_comprehensive_summary.png"):
                    st.markdown("#### Comprehensive Research Impact Summary")
                    st.image("research_comprehensive_summary.png", use_column_width=True)
            
            else:  # Generate All
                if st.button("🎨 Generate All Visualizations", type="primary", use_container_width=True):
                    with st.spinner("Generating all research visualizations and graphics..."):
                        # Generate research visualizations
                        from generate_research_visualizations import create_all_research_visualizations
                        create_all_research_visualizations()
                        
                        # Also generate presentation graphics
                        from generate_presentation_graphics import generate_all_graphics
                        generate_all_graphics()
                        
                        st.success("✅ All visualizations generated!")
                        st.balloons()
                        
                        st.markdown("""
                        ### 📁 Files Created:
                        
                        **Research Visualizations:**
                        - `research_quality_comparison.png`
                        - `research_performance_analysis.png` 
                        - `research_workflow_efficiency.png`
                        - `research_feedback_quality.png`
                        - `research_comprehensive_summary.png`
                        
                        **Presentation Graphics:**
                        - `performance_distribution.png`
                        - `ai_impact_comparison.png`
                        - `evaluation_timeline.png`
                        - `rubric_scores_radar.png`
                        - `subject_area_distribution.png`
                        - `evaluation_workflow.png`
                        """)
        
        with tab2:
            st.markdown("### 5 Demo Scenarios Available")
            
            # Simple list of scenarios
            scenarios = [
                {"name": "Sarah Martinez", "subject": "Elementary Science", "performance": "Exemplary (91%)", "highlight": "Best practices demonstration"},
                {"name": "Michael Chen", "subject": "Secondary Math", "performance": "Progressing (55%)", "highlight": "Growth tracking example"},
                {"name": "Emily Johnson", "subject": "Special Education", "performance": "Highly Proficient (80%)", "highlight": "Inclusive practices"},
                {"name": "Jessica Taylor", "subject": "Kindergarten", "performance": "Exemplary (95%)", "highlight": "Early childhood excellence"},
                {"name": "David Park", "subject": "High School English", "performance": "Proficient (75%)", "highlight": "Technology integration"}
            ]
            
            for i, scenario in enumerate(scenarios):
                with st.expander(f"**{scenario['name']}** - {scenario['subject']}"):
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.markdown(f"**Performance:** {scenario['performance']}")
                        st.markdown(f"**Key Feature:** {scenario['highlight']}")
                    with col2:
                        st.info("Ready for demo")
        
        with tab3:
            st.markdown("### 📸 Export Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("💾 Save All Charts as PNG", type="primary", use_container_width=True):
                    with st.spinner("Generating high-resolution graphics..."):
                        from generate_presentation_graphics import generate_all_graphics
                        generate_all_graphics()
                        st.success("✅ Saved 6 PNG files to project directory!")
                        
                        st.markdown("""
                        **Files created:**
                        - performance_distribution.png
                        - ai_impact_comparison.png
                        - evaluation_timeline.png
                        - rubric_scores_radar.png
                        - subject_area_distribution.png
                        - evaluation_workflow.png
                        """)
            
            with col2:
                st.info("""
                **Tips for poster:**
                - Take screenshots of charts above
                - Or use the PNG files generated
                - All charts use UVU colors
                - 300 DPI resolution for printing
                """)
    
    else:
        # Show instructions if demo not ready
        st.markdown("---")
        st.info("👆 Click the **Generate Demo Data & Charts** button above to get started!")

def show_settings():
    """Settings and configuration"""
    st.header("⚙️ Settings")
    
    # OpenAI Configuration
    st.subheader("🤖 AI Configuration")
    
    # Show current model status
    current_model_display = openai_service.model if openai_service else os.getenv('OPENAI_MODEL', 'gpt-5-mini')
    st.info(f"🎯 Currently using model: **{current_model_display}**")
    
    # Initialize session state for API key
    if 'api_key' not in st.session_state:
        st.session_state.api_key = os.getenv('OPENAI_API_KEY', '')
    
    # Check current API key status
    current_key = os.getenv('OPENAI_API_KEY', '') or st.session_state.api_key
    
    if current_key:
        st.success("✅ API key is configured! AI features are available.")
        st.info(f"Using API key: {current_key[:7]}{'*' * 20}{current_key[-4:]}")
        
        if st.button("🔄 Update API Key"):
            st.session_state.show_api_input = True
    else:
        st.warning("⚠️ No API key configured. AI features are disabled.")
        st.session_state.show_api_input = True
    
    # Show API key input if needed
    if st.session_state.get('show_api_input', False) or not current_key:
        api_key = st.text_input(
            "OpenAI API Key",
            value=st.session_state.api_key,
            type="password",
            help="Your OpenAI API key for AI-powered features",
            placeholder="Enter your OpenAI API key"
        )
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("💾 Save API Key"):
                if api_key.strip():
                    os.environ['OPENAI_API_KEY'] = api_key.strip()
                    st.session_state.api_key = api_key.strip()
                    st.session_state.show_api_input = False
                    st.success("API key saved for this session!")
                    st.rerun()
                else:
                    st.error("Please enter a valid API key")
        
        with col2:
            if st.button("❌ Cancel"):
                st.session_state.show_api_input = False
                st.rerun()
        
        st.caption("💡 **Note**: API key is saved for this session only. Set OPENAI_API_KEY environment variable for permanent storage.")
    
    # Get current model from OpenAI service
    current_model = openai_service.model if openai_service else os.getenv('OPENAI_MODEL', 'gpt-5-mini')
    
    # Create model options with current model first
    model_options = ["gpt-5-mini", "gpt-5-nano", "gpt-5"]
    if current_model in model_options:
        model_options.remove(current_model)
        model_options.insert(0, current_model)
    
    model = st.selectbox(
        "OpenAI Model",
        model_options,
        index=0,  # Select the current model
        help=f"Currently using: {current_model}. Choose a different model to override for this session."
    )
    
    # Update model if changed
    if model != current_model and openai_service:
        if st.button("🔄 Apply Model Change"):
            openai_service.model = model
            st.success(f"Model updated to {model}")
            st.rerun()
    
    # AI Debugging (if there were recent errors)
    if 'last_ai_error' in st.session_state:
        with st.expander("🐛 Debug: Last AI Error", expanded=False):
            error_info = st.session_state.last_ai_error
            st.warning("Recent AI parsing error detected")
            st.write(f"**Timestamp:** {error_info.get('timestamp', 'Unknown')}")
            st.write(f"**Model:** {error_info.get('model', 'Unknown')}")
            st.write("**Parsing Errors:**")
            for err in error_info.get('errors', []):
                st.code(err)
            st.write("**AI Response Preview:**")
            st.code(error_info.get('response_preview', 'No preview available'))
            
            if st.button("Clear Debug Info"):
                del st.session_state.last_ai_error
                st.rerun()
    
    # App Configuration
    st.subheader("📱 Application Settings")
    
    theme = st.selectbox("Theme", ["light", "dark", "auto"])
    
    # Data Management
    st.subheader("💾 Data Management")
    
    if st.button("Clear All Data"):
        if st.checkbox("I understand this will delete all evaluations"):
            # Implementation needed
            st.warning("Clear data functionality not implemented yet")
    
    # About
    st.subheader("ℹ️ About AI-STER")
    st.write("""
    **AI-STER** (AI-powered Student Teaching Evaluation Rubric) is a comprehensive evaluation system 
    for student teachers, aligned with USBE standards and enhanced with AI capabilities.
    
    - **Version:** 1.0.0
    - **USBE Compliance:** July 2024 standards
    - **AI Features:** OpenAI-powered evaluation assistance
    """)

def get_level_name(level) -> str:
    """Get the name for a scoring level"""
    if level == "not_observed":
        return "Not Observed"
    elif isinstance(level, int):
        names = {
            0: "Does not demonstrate",
            1: "Approaching", 
            2: "Demonstrates",
            3: "Exceeds"
        }
        return names.get(level, "Unknown")
    else:
        return "Unknown"

def get_disposition_level_name(level: int) -> str:
    """Get the name for a disposition scoring level"""
    names = {
        1: "Does not demonstrate disposition",
        2: "Is approaching disposition at expected level", 
        3: "Demonstrates disposition at expected level",
        4: "Exceeds expectations"
    }
    return names.get(level, "Unknown")

if __name__ == "__main__":
    main() 